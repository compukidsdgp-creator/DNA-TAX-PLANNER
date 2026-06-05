# ─────────────────────────────────────────────────────────────────────────────
#  DNA Accountants Pty Ltd — Tax Optimisation Report Generator  v3.0
#  FY2025-26  |  Sydney, Australia
#  Enhancements: 2-yr comparison, tax bracket viz, income-split, rental CF,
#  exec summary, YoY table, audit risk flags, super modeller, CGT harvest,
#  batch processing, responsive wizard UI, mobile layout
# ─────────────────────────────────────────────────────────────────────────────
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
import io, os, json, re, base64, pathlib, copy, warnings, zipfile, tempfile
from datetime import datetime
import requests
warnings.filterwarnings("ignore")

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="DNA Accountants | Tax Report v3",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Brand ─────────────────────────────────────────────────────────────────────
DNA_BLUE  = "#1B3A6B"
DNA_GOLD  = "#C9A84C"
DNA_GREEN = "#27AE60"
DNA_RED   = "#C0392B"
COMPANY   = "DNA Accountants Pty Ltd"

import base64, pathlib

def _load_logo_b64():
    """Try multiple paths to find the logo — robust for Streamlit Cloud and local."""
    candidates = [
        pathlib.Path(__file__).parent / "dna_logo.png",
        pathlib.Path("dna_logo.png"),
        pathlib.Path(__file__).resolve().parent / "dna_logo.png",
    ]
    for p in candidates:
        try:
            if p.exists():
                return base64.b64encode(p.read_bytes()).decode()
        except Exception:
            pass
    return ""

LOGO_B64 = _load_logo_b64()
_LOGO_PATH = pathlib.Path(__file__).parent / "dna_logo.png"

# ── Wizard step state ─────────────────────────────────────────────────────────
if "step" not in st.session_state:
    st.session_state.step = 1          # 1=Upload  2=Review  3=AI Advice  4=Download
if "f1"   not in st.session_state: st.session_state.f1   = None
if "f2"   not in st.session_state: st.session_state.f2   = None
if "meta" not in st.session_state: st.session_state.meta = {}
if "advice" not in st.session_state: st.session_state.advice = []
if "batch_results" not in st.session_state: st.session_state.batch_results = []
if "use_ai" not in st.session_state: st.session_state.use_ai = True

# ─────────────────────────────────────────────────────────────────────────────
#  CSS
# ─────────────────────────────────────────────────────────────────────────────
st.markdown(f"""
<style>
  /* ── global ── */
  [data-testid="stAppViewContainer"] {{ background:#F4F7FC; }}
  [data-testid="stSidebar"] {{ background:{DNA_BLUE}; }}
  [data-testid="stSidebar"] * {{ color:white !important; }}
  /* ── sidebar nav buttons (wizard) ── */
  [data-testid="stSidebar"] .stButton>button {{
      background:rgba(201,168,76,0.18) !important;
      color:{DNA_GOLD} !important;
      border:1px solid rgba(201,168,76,0.50) !important;
      border-radius:8px !important;
      font-weight:600 !important;
      font-size:13px !important;
      transition:all .2s !important;
  }}
  [data-testid="stSidebar"] .stButton>button:hover {{
      background:{DNA_GOLD} !important;
      color:{DNA_BLUE} !important;
      border-color:{DNA_GOLD} !important;
  }}

  /* ── sidebar download buttons ── */
  [data-testid="stSidebar"] .stDownloadButton>button {{
      background:{DNA_GOLD} !important; color:{DNA_BLUE} !important;
      font-weight:700 !important;
  }}

  /* ── header banner ── */
  .dna-header {{
      background:linear-gradient(135deg,{DNA_BLUE} 0%,#1e4d8c 60%,#163d72 100%);
      color:white; padding:22px 32px; border-radius:14px; margin-bottom:4px;
      display:flex; align-items:center; gap:28px;
      box-shadow:0 4px 18px rgba(27,58,107,.28);
  }}
  .dna-header img {{ height:76px; width:auto; display:block; }}
  .dna-header-div {{ width:2px; height:76px; background:rgba(201,168,76,.55); flex-shrink:0; }}
  .dna-header-text h1 {{ margin:0 0 5px; font-size:24px; font-weight:800; color:white; }}
  .dna-header-text p  {{ margin:0 0 3px; font-size:12px; }}
  .gold-line {{ height:4px; background:{DNA_GOLD}; border-radius:2px; margin:0 0 18px; }}

  /* ── wizard progress bar ── */
  .wizard-bar {{
      display:flex; gap:2px; margin-bottom:22px; border-radius:10px;
      overflow:hidden; box-shadow:0 3px 10px rgba(27,58,107,.20);
      background:{DNA_BLUE};
  }}
  .wz-step {{
      flex:1; padding:13px 8px; text-align:center; font-size:12px;
      font-weight:600; cursor:pointer; transition:all .2s;
      background:rgba(255,255,255,0.10); color:rgba(255,255,255,0.65);
      border-right:1px solid rgba(255,255,255,0.12);
  }}
  .wz-step:last-child {{ border-right:none; }}
  .wz-step:hover {{
      background:rgba(201,168,76,0.30); color:white;
  }}
  .wz-step.active {{
      background:{DNA_GOLD}; color:white;
      box-shadow:inset 0 -3px 0 rgba(0,0,0,0.20);
  }}
  .wz-step.done {{
      background:rgba(39,174,96,0.75); color:white;
  }}
  .wz-step.done:hover {{ background:rgba(39,174,96,0.90); }}
  .wz-step .step-num {{ font-size:18px; font-weight:800; display:block; }}

  /* ── KPI cards ── */
  .kpi-card {{
      background:white; border-radius:10px; padding:18px 20px;
      border-left:5px solid {DNA_BLUE}; box-shadow:0 2px 8px rgba(0,0,0,.07);
  }}
  .kpi-card .label {{ font-size:11px; color:#666; text-transform:uppercase; letter-spacing:.5px; }}
  .kpi-card .value {{ font-size:24px; font-weight:700; color:{DNA_BLUE}; }}
  .kpi-card .delta {{ font-size:12px; margin-top:2px; }}
  .delta-up   {{ color:{DNA_GREEN}; }}
  .delta-down {{ color:{DNA_RED}; }}

  /* ── section heading ── */
  .sec-head {{
      font-size:17px; font-weight:700; color:{DNA_BLUE};
      border-bottom:2px solid {DNA_GOLD}; padding-bottom:5px; margin:22px 0 14px;
  }}

  /* ── advice cards ── */
  .adv-card {{
      background:white; border-radius:10px; padding:16px 20px;
      border-left:5px solid {DNA_GOLD}; box-shadow:0 2px 8px rgba(0,0,0,.07);
      margin-bottom:12px;
  }}
  .adv-card .atitle {{ font-size:14px; font-weight:700; color:{DNA_BLUE}; margin-bottom:5px; }}
  .adv-card .abody  {{ font-size:13px; color:#444; line-height:1.6; }}

  /* ── risk flag ── */
  .risk-high   {{ background:#FFF0F0; border-left:5px solid {DNA_RED}; border-radius:8px;
                  padding:10px 14px; margin-bottom:8px; font-size:13px; }}
  .risk-medium {{ background:#FFFBF0; border-left:5px solid {DNA_GOLD}; border-radius:8px;
                  padding:10px 14px; margin-bottom:8px; font-size:13px; }}
  .risk-low    {{ background:#F0FFF4; border-left:5px solid {DNA_GREEN}; border-radius:8px;
                  padding:10px 14px; margin-bottom:8px; font-size:13px; }}

  /* ── exec summary box ── */
  .exec-box {{
      background:white; border:2px solid {DNA_BLUE}; border-radius:12px;
      padding:24px 28px; margin-bottom:20px;
  }}
  .exec-box h2 {{ color:{DNA_BLUE}; margin:0 0 12px; font-size:17px; }}

  /* ── download buttons ── */
  .stDownloadButton>button {{
      background:{DNA_BLUE} !important; color:white !important;
      border-radius:8px !important; font-weight:600 !important;
      width:100%;
  }}
  .stDownloadButton>button:hover {{ background:#2A5298 !important; }}

  /* ── responsive ── */
  @media (max-width: 768px) {{
      .dna-header {{ flex-direction:column; gap:14px; padding:16px; }}
      .dna-header img {{ height:54px; }}
      .dna-header-div {{ display:none; }}
      .dna-header-text h1 {{ font-size:18px; }}
      .kpi-card .value {{ font-size:18px; }}
      .wz-step {{ font-size:11px; padding:8px 4px; }}
      .wz-step .step-num {{ font-size:14px; }}
  }}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
#  HEADER
# ─────────────────────────────────────────────────────────────────────────────
_logo = (f'<img src="data:image/png;base64,{LOGO_B64}" alt="DNA">' if LOGO_B64
         else '<span style="font-size:52px">🏛️</span>')
st.markdown(f"""
<div class="dna-header">
  <div>{_logo}</div>
  <div class="dna-header-div"></div>
  <div class="dna-header-text">
    <h1>DNA Accountants Pty Ltd</h1>
    <p style="color:#BDD4F0;">Chartered Accountants &bull; Sydney, Australia</p>
    <p style="color:#D0E2F5;font-size:11px;">
        AI-Powered Tax Optimisation Report &bull; FY2024–25 &amp; FY2025–26 &bull; v3.0
    </p>
  </div>
</div>
<div class="gold-line"></div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
#  SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    _sb_logo = LOGO_B64 or _load_logo_b64()
    if _sb_logo:
        st.markdown(
            f'<div style="text-align:center;padding:16px 0 6px;">' +
            f'<img src="data:image/png;base64,{_sb_logo}" ' +
            f'style="height:64px;width:auto;max-width:160px;display:block;margin:0 auto;">' +
            f'</div>', unsafe_allow_html=True)
    st.markdown('<div style="text-align:center;font-size:10px;opacity:.7;'
                'padding-bottom:10px;">Chartered Accountants</div>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("## ⚙️ Settings")

    # ── AI Mode toggle ────────────────────────────────────────────────────────
    use_ai = st.toggle("🤖 Use AI-Powered Advice", value=True,
                       help="ON = Gemini AI generates advice. OFF = rule-based advice (no API needed).")

    if use_ai:
        # Try Streamlit Cloud secrets first, fall back to manual entry
        _secret_key = ""
        try:
            _secret_key = st.secrets.get("GEMINI_API_KEY", "")
        except Exception:
            pass

        if _secret_key:
            gemini_key = _secret_key
            st.markdown(
                '<div style="background:rgba(39,174,96,0.18);border-left:3px solid #27AE60;'
                'border-radius:6px;padding:8px 10px;font-size:11px;color:#E0F0E8;margin-top:6px;">'
                '🔒 <b>API key loaded from server secrets</b><br>'
                'Key is private &amp; secure — not visible to anyone.</div>',
                unsafe_allow_html=True)
        else:
            gemini_key = st.text_input(
                "🔑 Enter Gemini API Key",
                type="password",
                help="Paste your Google Gemini Flash 2.5 key. Get one free at aistudio.google.com",
                placeholder="AIza...")
            if gemini_key:
                st.markdown(
                    '<div style="background:rgba(39,174,96,0.18);border-left:3px solid #27AE60;'
                    'border-radius:6px;padding:6px 10px;font-size:11px;color:#E0F0E8;">'
                    '✅ Key entered — not stored anywhere.</div>',
                    unsafe_allow_html=True)
            else:
                st.markdown(
                    '<div style="background:rgba(201,168,76,0.18);border-left:3px solid #C9A84C;'
                    'border-radius:6px;padding:6px 10px;font-size:11px;color:#F5EED0;">'
                    '💡 No key yet — will use rule-based advice.</div>',
                    unsafe_allow_html=True)
    else:
        gemini_key = ""
        st.markdown(
            '<div style="background:rgba(255,255,255,0.10);border-left:3px solid #C9A84C;'
            'border-radius:6px;padding:8px 10px;font-size:11px;color:#D0E2F5;margin-top:4px;">'
            '📋 <b>Rule-based mode</b><br>'
            "Using DNA's built-in ATO tax optimisation logic.<br>"
            'No API key required.</div>',
            unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 🧙 Wizard Navigation")
    steps = ["1 · Upload", "2 · Review", "3 · AI Advice", "4 · Download"]
    for i, label in enumerate(steps, 1):
        icon = "✅" if st.session_state.step > i else ("▶️" if st.session_state.step == i else "⬜")
        if st.button(f"{icon} {label}", key=f"nav_{i}", use_container_width=True):
            if i <= st.session_state.step:
                st.session_state.step = i
                st.rerun()
    st.markdown("---")
    st.markdown("### 📋 CSV Template")

    # Two-year sample CSV
    sample_rows = [
        ("client_name",                        "Michael Thompson",  "Michael Thompson"),
        ("fy_year",                             "FY2024-25",        "FY2025-26"),
        ("salary",                              "125000",           "135000"),
        ("pension_income",                      "0",                "0"),
        ("allowances",                          "3800",             "4200"),
        ("other_employment_income",             "1500",             "1800"),
        ("rental_income",                       "24000",            "26400"),
        ("rental_expenses",                     "7200",             "7800"),
        ("rental_loan_interest",                "10800",            "11200"),
        ("rental_depreciation",                 "4200",             "4500"),
        ("rental_property_value",               "680000",           "710000"),
        ("rental_loan_balance",                 "420000",           "410000"),
        ("business_income",                     "0",                "0"),
        ("business_expenses",                   "0",                "0"),
        ("capital_gains",                       "8000",             "22000"),
        ("capital_gains_discount_eligible",     "yes",              "yes"),
        ("capital_losses_available",            "2000",             "1500"),
        ("unrealised_loss_assets",              "3500",             "4200"),
        ("other_income",                        "1400",             "1850"),
        ("payg_withholding",                    "29800",            "34200"),
        ("payg_instalments",                    "2200",             "3100"),
        ("super_contributions_concessional",    "5000",             "5500"),
        ("super_contributions_non_concessional","0",                "0"),
        ("super_balance",                       "145000",           "162000"),
        ("spouse_income",                       "48000",            "52000"),
        ("work_related_deductions",             "2800",             "3200"),
        ("self_education_deductions",           "900",              "1200"),
        ("charitable_donations",                "600",              "750"),
        ("income_protection_insurance",         "1500",             "1800"),
        ("other_deductions",                    "350",              "420"),
        ("low_income_tax_offset_eligible",      "no",               "no"),
        ("seniors_offset_eligible",             "no",               "no"),
        ("private_health_insurance",            "yes",              "yes"),
        ("has_hecs_debt",                       "yes",              "yes"),
        ("hecs_balance",                        "22000",            "18500"),
    ]
    sample_df = pd.DataFrame(sample_rows, columns=["field", "fy1_value", "fy2_value"])
    csv_template = sample_df.to_csv(index=False)
    st.download_button("⬇️ Download 2-Year CSV Template", csv_template,
                       "dna_tax_2yr_template.csv", "text/csv", use_container_width=True)
    st.markdown("---")
    st.markdown("**DNA Accountants Pty Ltd**  \nChartered Accountants  \n"
                "📍 110 Pitt Street, Sydney NSW 2000  \n📞 (02) 9064 4400")

# ─────────────────────────────────────────────────────────────────────────────
#  WIZARD PROGRESS BAR
# ─────────────────────────────────────────────────────────────────────────────
def render_wizard_bar():
    labels = ["📂 Upload & Parse", "📊 Review & Analysis", "🤖 AI Advice", "⬇️ Download Report"]
    html = '<div class="wizard-bar">'
    for i, lbl in enumerate(labels, 1):
        cls = "active" if st.session_state.step == i else ("done" if st.session_state.step > i else "wz-step")
        if cls == "active":
            cls = "wz-step active"
        elif cls == "done":
            cls = "wz-step done"
        else:
            cls = "wz-step"
        html += f'<div class="{cls}"><span class="step-num">{i}</span>{lbl}</div>'
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)

render_wizard_bar()

# ─────────────────────────────────────────────────────────────────────────────
#  HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def safe_float(val, default=0.0):
    try:
        return float(str(val).replace(',','').replace('$','').strip())
    except:
        return default

def safe_bool(val):
    """Handle Yes/No/yes/no/YES/NO/true/false/1/0."""
    if val is None: return False
    return str(val).strip().lower() in ('yes','true','1','y','✓','x')

def calc_tax(ti):
    if ti <= 0: return 0
    tax, prev = 0, 0
    for thresh, rate in [(18200,0),(45000,0.19),(135000,0.30),(190000,0.37),(float('inf'),0.45)]:  # Stage 3 tax cuts
        if ti <= thresh:
            return round(tax + (ti - prev) * rate, 2)
        tax += (thresh - prev) * rate
        prev = thresh
    return round(tax, 2)

def calc_medicare(ti):
    return 0 if ti <= 26000 else round(ti * 0.02, 2)

def lito(ti):
    if ti <= 37500:  return 700
    if ti <= 45000:  return max(0, 700  - (ti - 37500) * 0.05)
    if ti <= 66667:  return max(0, 325  - (ti - 45000) * 0.015)
    return 0

SUPER_CAP = 30000

def compute(d, yr_suffix=""):
    """Compute all financials from a dict. yr_suffix = '' or '_py' """
    def g(k): return safe_float(d.get(k, 0))
    def b(k): return safe_bool(d.get(k, 'no'))

    f = {}
    f['salary']       = g('salary')
    f['pension']       = g('pension_income')
    f['allowances']    = g('allowances')
    f['other_emp']     = g('other_employment_income')
    f['rental_income'] = g('rental_income')
    f['rental_exp']    = g('rental_expenses')
    f['rental_int']    = g('rental_loan_interest')
    f['rental_dep']    = g('rental_depreciation')
    f['prop_value']    = g('rental_property_value')
    f['loan_balance']  = g('rental_loan_balance')
    f['biz_income']    = g('business_income')
    f['biz_exp']       = g('business_expenses')
    f['cgt_gross']     = g('capital_gains')
    f['cgt_discount']  = b('capital_gains_discount_eligible')
    f['cap_losses']    = g('capital_losses_available')
    f['unrealised_loss']= g('unrealised_loss_assets')
    f['other_income']  = g('other_income')
    f['payg_wh']       = g('payg_withholding')
    f['payg_inst']     = g('payg_instalments')
    # ── Super: correctly separate employer SG from personal deductible ──────
    # super_contributions_concessional = employer SG + any salary sacrifice (display only)
    # super_personal_deductible = personal contributions lodged under s290-150 ITAA97
    #   → ONLY this amount reduces taxable income on the personal tax return
    f['super_cc']     = g('super_contributions_concessional')  # total CC for display
    f['super_ded']    = g('super_personal_deductible')         # only this deducted from taxable
    f['super_ncc']    = g('super_contributions_non_concessional')
    f['super_bal']    = g('super_balance')
    f['franking_credits'] = g('franking_credits')              # dividend imputation offset
    f['spouse_inc']   = g('spouse_income')
    f['work_ded']     = g('work_related_deductions')
    f['self_ed']      = g('self_education_deductions')
    f['charity']      = g('charitable_donations')
    f['income_prot']  = g('income_protection_insurance')
    f['other_ded']    = g('other_deductions')
    f['lito_elig']    = b('low_income_tax_offset_eligible')
    f['seniors']      = b('seniors_offset_eligible')
    f['has_hecs']     = b('has_hecs_debt')
    f['hecs_bal']     = g('hecs_balance')
    f['priv_health']  = b('private_health_insurance')
    f['phi_excess']      = g('phi_excess_private_health')
    f['tax_adjustments'] = g('tax_adjustments')  # manual adj to reconcile with lodged return (negative = reduces tax)

    # ── Concessional super cap — correct calculation ────────────────────────
    # super_cc from CSV = RESC (reportable salary sacrifice, IT2 on ITR)
    # Mandatory SG (11.5% FY25/26) is never on the ITR — calculate it
    SG_RATE          = 0.115
    f['employer_sg'] = round(f['salary'] * SG_RATE, 2)   # mandatory non-reportable SG
    f['resc']        = f['super_cc']                       # reportable employer super (IT2)
    f['total_cc']    = f['employer_sg'] + f['resc']        # true total concessional
    _effective_cc    = f['total_cc']
    f['super_sg_estimated']  = False
    f['super_cap_exceeded']  = f['total_cc'] > SUPER_CAP
    f['super_excess']        = max(0, f['total_cc'] - SUPER_CAP)
    f['super_cc_display']    = f['total_cc']   # show true total in charts

    f['rental_net'] = f['rental_income'] - f['rental_exp'] - f['rental_int'] - f['rental_dep']
    f['biz_net']    = f['biz_income'] - f['biz_exp']
    cgt_net = max(0, f['cgt_gross'] - f['cap_losses'])
    f['cgt_taxable']= cgt_net * 0.5 if f['cgt_discount'] else cgt_net

    f['gross_income'] = (f['salary'] + f['pension'] + f['allowances'] + f['other_emp'] +
                         f['rental_income'] + f['biz_income'] + f['cgt_gross'] + f['other_income'])
    f['total_ded']    = (f['super_ded'] + f['work_ded'] + f['self_ed'] +
                         f['charity'] + f['income_prot'] + f['other_ded'])
    f['total_exp_ded']= (f['rental_exp'] + f['rental_int'] + f['rental_dep'] +
                         f['biz_exp'] + f['total_ded'])
    f['taxable_income'] = max(0,
        f['salary'] + f['pension'] + f['allowances'] + f['other_emp'] +
        f['rental_net'] + f['biz_net'] + f['cgt_taxable'] + f['other_income'] -
        f['super_ded'] -           # ONLY personal deductible super (s290-150 notice)
        f['work_ded'] - f['self_ed'] -
        f['charity'] - f['income_prot'] - f['other_ded']
    )

    it  = calc_tax(f['taxable_income'])
    med = calc_medicare(f['taxable_income'])
    lo  = lito(f['taxable_income']) if f['lito_elig'] else 0
    so  = 2230 if f['seniors'] else 0

    f['income_tax']   = it
    f['medicare']     = med
    f['lito']         = lo
    f['seniors_off']  = so
    phi = f.get('phi_excess', 0)
    adj = f.get('tax_adjustments', 0)
    f['gross_tax']    = max(0, it + med - lo - so - f['franking_credits'] + phi + adj)
    f['credits']      = f['payg_wh'] + f['payg_inst']
    f['tax_payable']  = f['gross_tax'] - f['credits']
    f['refund']       = max(0, -f['tax_payable'])
    f['owing']        = max(0,  f['tax_payable'])
    f['eff_rate']     = round(f['gross_tax'] / f['taxable_income'] * 100, 2) if f['taxable_income'] > 0 else 0

    # ATO-verified override (from lodged return / Xero)
    # ── ATO verified figures from lodged return ──────────────────
    ato_ti    = g('ato_taxable_income')
    ato_it    = g('ato_income_tax')
    ato_med   = g('ato_medicare_levy')
    ato_off   = g('ato_other_offsets')
    ato_gross = g('ato_gross_tax')
    ato_tax   = g('ato_tax_payable')     # positive=owing, negative=refund
    f['ato_verified']        = (ato_ti > 0 or ato_tax != 0)
    f['ato_taxable_income']  = ato_ti
    f['ato_income_tax']      = ato_it    if ato_it    > 0 else f['income_tax']
    f['ato_medicare_levy']   = ato_med   if ato_med   > 0 else f['medicare']
    f['ato_other_offsets']   = ato_off
    f['ato_gross_tax']       = ato_gross if ato_gross > 0 else f['gross_tax']
    f['ato_tax_payable_raw'] = ato_tax
    f['display_taxable']     = ato_ti    if ato_ti    > 0 else f['taxable_income']
    f['display_income_tax']  = ato_it    if ato_it    > 0 else f['income_tax']
    f['display_medicare']    = ato_med   if ato_med   > 0 else f['medicare']
    f['display_gross_tax']   = ato_gross if ato_gross > 0 else f['gross_tax']
    f['display_refund']      = max(0, -ato_tax) if f['ato_verified'] else f['refund']
    f['display_owing']       = max(0,  ato_tax) if f['ato_verified'] else f['owing']
    f['display_eff_rate']    = round(ato_gross / ato_ti * 100, 2) if (ato_ti > 0 and ato_gross > 0) else f['eff_rate']
    f['super_headroom'] = max(0, SUPER_CAP - _effective_cc)
    f['super_personal'] = f['super_headroom']

    # Rental cash-flow
    f['rental_cf']    = f['rental_income'] - f['rental_exp'] - f['rental_int']
    f['rental_tax_benefit'] = abs(min(0, f['rental_net'])) * (f['eff_rate'] / 100) if f['rental_net'] < 0 else 0

    # Spouse split saving estimate
    combined   = f['taxable_income'] + f['spouse_inc']
    split_each = combined / 2
    tax_current= f['gross_tax'] + calc_tax(f['spouse_inc'])
    tax_split  = calc_tax(split_each) * 2
    f['spouse_split_saving'] = max(0, tax_current - tax_split)

    # Super bracket saving
    brackets = [(18200,0),(45000,0.19),(135000,0.30),(190000,0.37),(float('inf'),0.45)]  # Stage 3
    f['marginal_rate'] = 0
    for thresh, rate in brackets:
        if f['taxable_income'] <= thresh:
            f['marginal_rate'] = rate
            break
    super_saving_per_k = (f['marginal_rate'] - 0.15) * 1000
    f['super_saving_per_k'] = max(0, super_saving_per_k)
    f['max_super_saving']   = max(0, f['super_headroom'] * (f['marginal_rate'] - 0.15))

    return f

def parse_two_year_csv(uploaded_file):
    df = pd.read_csv(uploaded_file)
    # Drop comment/section-header rows (field starts with #)
    df = df[~df['field'].astype(str).str.startswith('#')]
    if 'field' in df.columns:
        if 'fy1_value' in df.columns and 'fy2_value' in df.columns:
            d1, d2 = {}, {}
            for _, row in df.iterrows():
                k = str(row['field']).strip()
                d1[k] = str(row['fy1_value']).strip()
                d2[k] = str(row['fy2_value']).strip()
            return d1, d2
        elif 'value' in df.columns:
            d = {}
            for _, row in df.iterrows():
                d[str(row['field']).strip()] = str(row['value']).strip()
            return d, d   # single year — duplicate
    return {}, {}

# ─────────────────────────────────────────────────────────────────────────────
#  CHARTS
# ─────────────────────────────────────────────────────────────────────────────
def chart_yoy_comparison(f1, f2, fy1, fy2):
    """Side-by-side bar chart for key metrics across two years."""
    metrics = ['gross_income','taxable_income','gross_tax','total_ded']
    labels  = ['Gross Income','Taxable Income','Tax Payable','Total Deductions']
    v1 = [f1[m] for m in metrics]
    v2 = [f2[m] for m in metrics]
    x  = np.arange(len(labels)); w = 0.35

    fig, ax = plt.subplots(figsize=(9, 4.5))
    b1 = ax.bar(x - w/2, v1, w, label=fy1, color=DNA_BLUE,  edgecolor='white')
    b2 = ax.bar(x + w/2, v2, w, label=fy2, color=DNA_GOLD,  edgecolor='white')
    for bar in list(b1)+list(b2):
        h = bar.get_height()
        ax.text(bar.get_x()+bar.get_width()/2, h+max(max(v1),max(v2))*0.01,
                f"${h/1000:.0f}k", ha='center', fontsize=9, fontweight='bold')
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=10)
    ax.set_ylabel("AUD $"); ax.legend(fontsize=10)
    ax.set_title("Year-on-Year Key Metrics", fontsize=13, fontweight='bold', color=DNA_BLUE)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    fig.patch.set_facecolor('white'); plt.tight_layout()
    return fig

def chart_income_breakdown(f, fy):
    items = {'Salary':f['salary'],'Pension':f['pension'],'Allowances':f['allowances'],
             'Other Emp.':f['other_emp'],'Rental':f['rental_income'],
             'Business':f['biz_income'],'Capital Gains':f['cgt_gross'],'Other':f['other_income']}
    items = {k:v for k,v in items.items() if v > 0}
    if not items: return None
    fig, ax = plt.subplots(figsize=(7,4))
    colors = [DNA_BLUE,'#C9A84C','#3A7ABF','#E8B84B','#5B9BD5','#F0C060','#2E5F99','#D4A030']
    bars = ax.barh(list(items.keys()), list(items.values()), color=colors[:len(items)], edgecolor='white', height=0.6)
    mx = max(items.values())
    for bar,val in zip(bars, items.values()):
        ax.text(bar.get_width()+mx*.01, bar.get_y()+bar.get_height()/2,
                f"${val:,.0f}", va='center', fontsize=9)
    ax.set_title(f"Income Breakdown — {fy}", fontsize=12, fontweight='bold', color=DNA_BLUE)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    fig.patch.set_facecolor('white'); plt.tight_layout(); return fig

def chart_expense_breakdown(f, fy):
    items = {'Rental Expenses':f['rental_exp'],'Rental Interest':f['rental_int'],
             'Rental Depreciation':f['rental_dep'],'Business Exp.':f['biz_exp'],
             'Super (Concessional)':f['super_cc_display'],'Work-Related':f['work_ded'],
             'Self-Education':f['self_ed'],'Donations':f['charity'],
             'Income Protection':f['income_prot'],'Other':f['other_ded']}
    items = {k:v for k,v in items.items() if v > 0}
    if not items: return None
    fig, ax = plt.subplots(figsize=(7,4))
    colors = ["#C0392B","#E74C3C","#E67E22","#F39C12","#27AE60","#2ECC71","#16A085","#1ABC9C","#8E44AD","#9B59B6"]
    bars = ax.barh(list(items.keys()), list(items.values()), color=colors[:len(items)], edgecolor='white', height=0.6)
    mx = max(items.values())
    for bar,val in zip(bars, items.values()):
        ax.text(bar.get_width()+mx*.01, bar.get_y()+bar.get_height()/2,
                f"${val:,.0f}", va='center', fontsize=9)
    ax.set_title(f"Expenses & Deductions — {fy}", fontsize=12, fontweight='bold', color="#C0392B")
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    fig.patch.set_facecolor('white'); plt.tight_layout(); return fig

def chart_tax_bracket(f, fy):
    """Tax bracket visualiser — colour-coded income gauge."""
    brackets = [(0,18200,'Tax-Free','#27AE60'),
                (18200,45000,'19%','#F1C40F'),
                (45000,120000,'32.5%','#E67E22'),
                (120000,180000,'37%','#E74C3C'),
                (180000,300000,'45%','#C0392B')]
    ti = f['taxable_income']
    cap = max(300000, ti * 1.1)

    fig, ax = plt.subplots(figsize=(10, 3.2))
    fig.subplots_adjust(top=0.72, bottom=0.18)   # leave room for title above & axis below

    for low, high, label, color in brackets:
        w = min(high, cap) - low
        ax.barh(0, w, left=low, height=0.55, color=color, edgecolor='white', linewidth=1.5)
        mid = (low + min(high, cap)) / 2
        ax.text(mid, 0, label, ha='center', va='center', fontsize=9,
                fontweight='bold', color='white')

    # Client income marker — line only inside bar, label ABOVE the bar
    ax.axvline(ti, color=DNA_BLUE, linewidth=2.5, zorder=5)
    # Annotate above the bar (y > 0.3) so it never overlaps the title
    ax.annotate(f"${ti:,.0f}",
                xy=(ti, 0.28), xytext=(ti, 0.42),
                fontsize=9, fontweight='bold', color=DNA_BLUE,
                ha='center', va='bottom',
                arrowprops=dict(arrowstyle='-', color=DNA_BLUE, lw=1.2))

    ax.set_xlim(0, cap)
    ax.set_ylim(-0.35, 0.55)
    ax.set_yticks([])
    ax.set_xlabel("Taxable Income (AUD $)", fontsize=10)

    # Title as two lines: company name bold + report subtitle — placed well above the bar
    fig.text(0.5, 0.93,
             f"Tax Bracket Position — {fy}",
             ha='center', va='top', fontsize=11, fontweight='bold', color=DNA_BLUE)
    fig.text(0.5, 0.83,
             f"Marginal Rate: {int(f['marginal_rate']*100)}%   |   Effective Rate: {f['eff_rate']}%",
             ha='center', va='top', fontsize=10, color=DNA_BLUE)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    fig.patch.set_facecolor('white')
    return fig

def chart_rental_cashflow(f, fy):
    """Rental property cash-flow vs tax benefit waterfall."""
    items = {
        'Rental Income': f['rental_income'],
        'Less: Expenses': -f['rental_exp'],
        'Less: Interest': -f['rental_int'],
        'Cash Flow':       f['rental_cf'],
        'Less: Depreciation': -f['rental_dep'],
        'Net Rental':      f['rental_net'],
        'Tax Benefit':     f['rental_tax_benefit'],
    }
    labels = list(items.keys())
    values = list(items.values())
    colors = [DNA_GREEN if v >= 0 else DNA_RED for v in values]
    colors[3] = DNA_BLUE   # cash flow
    colors[5] = '#8E44AD'  # net rental
    colors[6] = DNA_GOLD   # tax benefit

    fig, ax = plt.subplots(figsize=(9, 4))
    bars = ax.bar(labels, values, color=colors, edgecolor='white', width=0.6)
    for bar, val in zip(bars, values):
        ypos = bar.get_height() if val >= 0 else bar.get_height() + val
        ax.text(bar.get_x()+bar.get_width()/2,
                (val + 500 if val >= 0 else val - 500),
                f"${val:,.0f}", ha='center', va='bottom' if val>=0 else 'top',
                fontsize=9, fontweight='bold')
    ax.axhline(0, color='black', linewidth=0.8)
    ax.set_title(f"Rental Property Cash Flow & Tax Analysis — {fy}", fontsize=12, fontweight='bold', color=DNA_BLUE)
    ax.set_ylabel("AUD $")
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    plt.xticks(rotation=15, ha='right', fontsize=9)
    fig.patch.set_facecolor('white'); plt.tight_layout(); return fig

def chart_super_modeller(f, fy):
    """Show tax saving at different additional super contribution levels."""
    headroom = f['super_headroom']
    if headroom <= 0:
        return None
    steps   = np.linspace(0, headroom, 20)
    savings = [max(0, s * (f['marginal_rate'] - 0.15)) for s in steps]

    fig, ax = plt.subplots(figsize=(8, 3.8))
    ax.fill_between(steps, savings, alpha=0.15, color=DNA_GOLD)
    ax.plot(steps, savings, color=DNA_GOLD, linewidth=2.5)
    ax.set_xlabel("Additional Concessional Super Contribution ($)", fontsize=10)
    ax.set_ylabel("Estimated Tax Saving ($)", fontsize=10)
    ax.set_title(f"Super Contribution Tax Saving Modeller — {fy}\n"
                 f"Marginal Rate {int(f['marginal_rate']*100)}% → Super Tax 15% = "
                 f"{int((f['marginal_rate']-0.15)*100)}¢ saving per $1 contributed",
                 fontsize=11, fontweight='bold', color=DNA_BLUE)

    # Mark headroom
    ax.axvline(headroom, color=DNA_RED, linestyle='--', linewidth=1.5)
    ax.text(headroom, max(savings)*0.5, f"  Cap limit\n  ${headroom:,.0f}",
            color=DNA_RED, fontsize=9)
    # Optimal point
    opt_saving = savings[-1]
    ax.scatter([headroom], [opt_saving], color=DNA_GOLD, s=80, zorder=5)
    ax.text(headroom*0.5, opt_saving*0.7, f"Max saving: ${opt_saving:,.0f}",
            fontsize=10, fontweight='bold', color=DNA_BLUE)

    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    fig.patch.set_facecolor('white'); plt.tight_layout(); return fig

def chart_cgt_harvest(f, fy):
    """CGT harvest opportunity chart."""
    cgt = f['cgt_gross']
    losses_avail = f['cap_losses']
    unrealised   = f['unrealised_loss']
    if cgt == 0 and unrealised == 0:
        return None

    tax_no_harvest   = calc_tax(f['taxable_income']) - calc_tax(
        max(0, f['taxable_income'] - (f['cgt_taxable'])))
    reduced_gain     = max(0, cgt - losses_avail - unrealised)
    reduced_taxable  = max(0, f['taxable_income'] - f['cgt_taxable'] +
                           (reduced_gain * 0.5 if f['cgt_discount'] else reduced_gain))
    tax_with_harvest = calc_tax(reduced_taxable) - calc_tax(
        max(0, reduced_taxable - (reduced_gain * 0.5 if f['cgt_discount'] else reduced_gain)))
    saving = max(0, tax_no_harvest - tax_with_harvest)

    cats   = ['Gross Capital Gain','Available Losses','Unrealised Losses\n(Harvestable)','Net Gain After\nHarvest']
    values = [cgt, losses_avail, unrealised, max(0, cgt - losses_avail - unrealised)]
    colors = [DNA_RED, DNA_GREEN, DNA_GOLD, DNA_BLUE]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4), gridspec_kw={'width_ratios':[3,2]})

    bars = ax1.bar(cats, values, color=colors, edgecolor='white', width=0.55)
    for bar, val in zip(bars, values):
        ax1.text(bar.get_x()+bar.get_width()/2, bar.get_height()+cgt*0.02,
                 f"${val:,.0f}", ha='center', fontsize=9, fontweight='bold')
    ax1.set_title(f"CGT Harvest Analysis — {fy}", fontsize=11, fontweight='bold', color=DNA_BLUE)
    ax1.set_ylabel("AUD $")
    ax1.spines['top'].set_visible(False); ax1.spines['right'].set_visible(False)

    summary = ['Tax Without\nHarvesting', 'Est. Tax With\nHarvesting', 'Potential\nSaving']
    svals   = [tax_no_harvest, tax_with_harvest, saving]
    scols   = [DNA_RED, DNA_GOLD, DNA_GREEN]
    bars2   = ax2.bar(summary, svals, color=scols, edgecolor='white', width=0.5)
    for bar, val in zip(bars2, svals):
        ax2.text(bar.get_x()+bar.get_width()/2, bar.get_height()+max(svals)*0.03,
                 f"${val:,.0f}", ha='center', fontsize=10, fontweight='bold')
    ax2.set_title("Tax Impact", fontsize=11, fontweight='bold', color=DNA_BLUE)
    ax2.spines['top'].set_visible(False); ax2.spines['right'].set_visible(False)

    fig.patch.set_facecolor('white'); plt.tight_layout(); return fig

def chart_spouse_split(f, fy):
    if f['spouse_inc'] == 0: return None
    ti, sp = f['taxable_income'], f['spouse_inc']
    split  = (ti + sp) / 2
    tax_now   = f['gross_tax'] + calc_tax(sp)
    tax_split = calc_tax(split) * 2
    saving    = max(0, tax_now - tax_split)

    labels = [f"Current\n(You ${ti/1000:.0f}k / Spouse ${sp/1000:.0f}k)",
              f"After Income Split\n(${split/1000:.0f}k each)"]
    values = [tax_now, tax_split]
    colors = [DNA_RED, DNA_GREEN]

    fig, ax = plt.subplots(figsize=(6, 3.5))
    bars = ax.bar(labels, values, color=colors, edgecolor='white', width=0.45)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+max(values)*0.01,
                f"${val:,.0f}", ha='center', fontsize=11, fontweight='bold')
    if saving > 0:
        ax.annotate(f"Potential saving: ${saving:,.0f}",
                    xy=(0.5, 0.92), xycoords='axes fraction', ha='center',
                    fontsize=11, fontweight='bold', color=DNA_GREEN)
    ax.set_title(f"Household Income Split Analysis — {fy}", fontsize=11, fontweight='bold', color=DNA_BLUE)
    ax.set_ylabel("Combined Tax (AUD $)")
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    fig.patch.set_facecolor('white'); plt.tight_layout(); return fig

# ─────────────────────────────────────────────────────────────────────────────
#  AUDIT RISK FLAGS
# ─────────────────────────────────────────────────────────────────────────────
def audit_risk_flags(f, fy):
    flags = []
    ti = f['taxable_income']
    if ti > 0:
        ded_ratio = f['total_ded'] / ti
        if ded_ratio > 0.25:
            flags.append(("High", f"Work-related & personal deductions are {ded_ratio*100:.0f}% of taxable income — "
                          "ATO benchmarks flag claims above ~15–20% for this income range. Ensure all receipts are retained."))
    if f['rental_net'] < -20000:
        flags.append(("High", f"Net rental loss of ${abs(f['rental_net']):,.0f} is substantial. "
                      "ATO scrutinises large negative gearing claims — ensure all rental expense records, "
                      "loan statements and depreciation schedules are current."))
    if f['rental_income'] > 0 and f['rental_dep'] == 0:
        flags.append(("Medium", "Rental income declared but no depreciation claimed. "
                      "A quantity surveyor report may unlock additional legitimate deductions."))
    if f['cgt_gross'] > 50000:
        flags.append(("High", f"Capital gain of ${f['cgt_gross']:,.0f} is significant. "
                      "Confirm CGT discount eligibility (12-month holding), cost-base calculations and any rollover relief."))
    if f['charity'] > 5000:
        flags.append(("Medium", f"Charitable donations of ${f['charity']:,.0f} — ATO may request receipts. "
                      "Ensure all donations are to DGR-registered organisations."))
    if f['biz_income'] > 0 and f['biz_exp'] / max(1, f['biz_income']) > 0.8:
        flags.append(("High", f"Business expense ratio is {f['biz_exp']/f['biz_income']*100:.0f}% — "
                      "very high margins may trigger PSI or Division 7A review."))
    if f['work_ded'] > 5000:
        flags.append(("Medium", f"Work-related deductions of ${f['work_ded']:,.0f} exceed ATO's average for most occupations. "
                      "Ensure home office diary, logbook and receipts are in order."))
    if not flags:
        flags.append(("Low", "No significant audit risk indicators detected based on the data provided. "
                      "Maintain good record-keeping as always."))
    return flags

# ─────────────────────────────────────────────────────────────────────────────
#  GEMINI AI ADVICE
# ─────────────────────────────────────────────────────────────────────────────
def get_gemini_advice(api_key, f1, f2, fy1, fy2, client_name):
    lines = [
        "You are a senior Australian tax advisor at DNA Accountants Pty Ltd, Sydney.",
        "Client: " + client_name,
        "",
        "TWO-YEAR FINANCIAL SUMMARY (all amounts in AUD, no dollar signs):",
        f"  {fy1} Gross Income:     {int(f1['gross_income'])}",
        f"  {fy2} Gross Income:     {int(f2['gross_income'])}",
        f"  {fy1} Taxable Income:   {int(f1['taxable_income'])}",
        f"  {fy2} Taxable Income:   {int(f2['taxable_income'])}",
        f"  {fy1} Tax Paid:         {int(f1['gross_tax'])}",
        f"  {fy2} Tax Paid:         {int(f2['gross_tax'])}",
        f"  {fy2} Effective Rate:   {f2['eff_rate']} percent",
        f"  {fy2} Marginal Rate:    {int(f2['marginal_rate']*100)} percent",
        f"  {fy2} Super Concessional: {int(f2['super_cc'])}",
        f"  {fy2} Super Cap Headroom: {int(f2['super_headroom'])}",
        f"  {fy2} Super Balance:    {int(f2['super_bal'])}",
        f"  {fy2} Rental Net:       {int(f2['rental_net'])}",
        f"  {fy2} Capital Gains:    {int(f2['cgt_gross'])}",
        f"  {fy2} Unrealised Losses: {int(f2['unrealised_loss'])}",
        f"  {fy2} Spouse Income:    {int(f2['spouse_inc'])}",
        f"  {fy2} HECS Balance:     {int(f2['hecs_bal'])}",
        f"  {fy2} Work Deductions:  {int(f2['work_ded'])}",
        "",
        "INCOME TREND: income grew from " + str(int(f1['gross_income'])) + " to " + str(int(f2['gross_income'])),
        "TAX TREND: tax grew from " + str(int(f1['gross_tax'])) + " to " + str(int(f2['gross_tax'])),
        "",
        "Provide exactly 5 practical, Australia-specific tax optimisation tips for FUTURE years.",
        "Incorporate the 2-year trend in your analysis — note if the situation is improving or worsening.",
        "Reference ATO rules, section numbers, or specific strategies.",
        "",
        "IMPORTANT: respond with ONLY a valid JSON array, no markdown fences, no extra text.",
        "Each of the 5 objects must have exactly these keys:",
        '  "title"            - max 8 words, no dollar signs or special chars',
        '  "detail"           - 2-3 sentences, use plain numbers like AUD 1200',
        '  "estimated_saving" - plain text like AUD 1200 to AUD 2400 per year',
        '  "priority"         - exactly High, Medium, or Low',
        "Do NOT use dollar signs, curly braces, backslashes or special chars in JSON strings.",
        "Start response with [ and end with ].",
    ]
    prompt = "\n".join(lines)
    url  = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.3, "maxOutputTokens": 2000,
                             "responseMimeType": "application/json"}
    }
    try:
        resp = requests.post(url + "?key=" + api_key,
                             json=payload, headers={"Content-Type":"application/json"}, timeout=45)
        resp.raise_for_status()
        raw  = resp.json()['candidates'][0]['content']['parts'][0]['text'].strip()
        raw  = re.sub(r'^```[a-zA-Z]*\s*','',raw); raw = re.sub(r'\s*```$','',raw).strip()
        raw  = raw[raw.find('['):raw.rfind(']')+1]
        raw  = raw.replace('\u2018',"'").replace('\u2019',"'")
        raw  = raw.replace('\u201c','"').replace('\u201d','"')
        raw  = raw.replace('\u2013','-').replace('\u2014','-')
        result = json.loads(raw)
        return [{"title":str(t.get("title","")),"detail":str(t.get("detail","")),
                 "estimated_saving":str(t.get("estimated_saving","Varies")),
                 "priority":str(t.get("priority","Medium"))} for t in result[:5]]
    except Exception as e:
        return [{"title":"AI Response Error","detail":str(e),"estimated_saving":"N/A","priority":"Low"}]

def rule_based_advice(f2, fy2):
    tips = []
    # Super cap exceeded warning — shown first if applicable
    if f2['super_cap_exceeded']:
        tips.append({
            'title': 'Concessional Super Cap Already Exceeded',
            'detail': (
                f'Total concessional contributions in {fy2} are '
                f'AUD {int(f2["total_cc"])} (mandatory SG AUD {int(f2["employer_sg"])} '
                f'plus reportable salary sacrifice AUD {int(f2["resc"])}), which exceeds '
                f'the AUD 30000 cap by AUD {int(f2["super_excess"])}. The ATO will '
                f'include the excess in assessable income at your marginal rate less a '
                f'15 percent tax offset. No additional concessional contributions '
                f'should be made this year.'
            ),
            'estimated_saving': 'Review urgently to avoid ATO excess notice',
            'priority': 'High'
        })
    if not f2['super_cap_exceeded'] and f2['super_headroom'] > 2000:
        s = round(f2['max_super_saving'], 0)
        tips.append({"title":"Maximise Concessional Super Contributions",
                     "detail":f"You have AUD {int(f2['super_headroom'])} remaining in your FY26 "
                               f"concessional cap (AUD 30000). Salary sacrificing this amount saves "
                               f"approximately AUD {int(s)} in tax as contributions are taxed at 15 percent "
                               f"vs your {int(f2['marginal_rate']*100)} percent marginal rate (s290-150 ITAA97).",
                     "estimated_saving":f"AUD {int(s)} per year","priority":"High"})
    if f2['rental_net'] < 0:
        tips.append({"title":"Optimise Negative Gearing Strategy",
                     "detail":f"Your rental property generates a net loss of AUD {int(abs(f2['rental_net']))} "
                               f"this year. Ensure all legitimate expenses are claimed under s8-1 ITAA97. "
                               f"Consider commissioning a depreciation schedule if not already in place.",
                     "estimated_saving":f"AUD {int(f2['rental_tax_benefit'])} per year","priority":"Medium"})
    if f2['cgt_gross'] > 0 and f2['unrealised_loss'] > 0:
        est = round(f2['unrealised_loss'] * f2['marginal_rate'] * 0.5, 0)
        tips.append({"title":"Harvest Capital Losses Before 30 June",
                     "detail":f"You have AUD {int(f2['unrealised_loss'])} in unrealised losses. "
                               f"Crystallising these before 30 June can offset your AUD {int(f2['cgt_gross'])} "
                               f"capital gain under s102-5 ITAA97, reducing CGT significantly.",
                     "estimated_saving":f"AUD {int(est)} per year","priority":"High"})
    if f2['spouse_split_saving'] > 500:
        tips.append({"title":"Consider Household Income Splitting",
                     "detail":f"Structuring income-producing assets jointly could save up to "
                               f"AUD {int(f2['spouse_split_saving'])} in combined household tax. "
                               f"Consider investing in your spouse's name or through a family trust.",
                     "estimated_saving":f"AUD {int(f2['spouse_split_saving'])} per year","priority":"Medium"})
    if f2['has_hecs'] and f2['hecs_bal'] > 0:
        tips.append({"title":"Plan HECS Repayment Band Timing",
                     "detail":f"With AUD {int(f2['hecs_bal'])} HECS balance, compulsory repayments increase "
                               f"at income thresholds. Salary packaging reportable fringe benefits or timing "
                               f"deductible expenses can keep repayment income below the next band.",
                     "estimated_saving":"AUD 400 to AUD 1200 per year","priority":"Low"})
    # Pool of distinct fallback tips — drawn in order until we have 5
    fallbacks = [
        {"title":"Maximise Work-Related Expense Claims",
         "detail":"Ensure you are claiming all eligible work-related expenses under s8-1 ITAA97 "
                  "including home office (67c/hr fixed rate), mobile phone work use, "
                  "professional development and tools. Maintain a diary and receipts.",
         "estimated_saving":"AUD 300 to AUD 1500 per year","priority":"Low"},

        {"title":"Review Income Protection Insurance Deductibility",
         "detail":f"Premiums paid for income protection insurance outside of super are fully "
                  f"tax-deductible under s8-1 ITAA97. At your {int(f2['marginal_rate']*100)} percent "
                  f"marginal rate, a AUD 2000 annual premium saves approximately AUD "
                  f"{int(2000 * f2['marginal_rate'])} in tax while protecting your income.",
         "estimated_saving":"AUD 400 to AUD 900 per year","priority":"Low"},

        {"title":"Prepay Deductible Expenses Before 30 June",
         "detail":"Under s82KZMF ITAA36, prepaying up to 12 months of deductible expenses before "
                  "30 June brings the deduction into the current year. This includes income protection "
                  "premiums, investment loan interest, professional memberships and subscriptions.",
         "estimated_saving":"AUD 500 to AUD 2000 per year","priority":"Medium"},

        {"title":"Consider a Testamentary Trust in Your Will",
         "detail":"A testamentary trust allows income from inherited assets to be distributed to "
                  "beneficiaries (including minor children) at adult tax rates, potentially saving "
                  "significant tax across generations. Consult your DNA advisor about estate planning.",
         "estimated_saving":"Long-term estate planning benefit","priority":"Low"},

        {"title":"Review Timing of Assessable Income",
         "detail":f"If you expect lower income in the next financial year, consider deferring "
                  f"assessable income (e.g. discretionary trust distributions, timing of asset sales) "
                  f"to reduce your current year tax at your {int(f2['marginal_rate']*100)} percent marginal rate.",
         "estimated_saving":"Situation-dependent","priority":"Medium"},
    ]

    # Add fallbacks that are not already covered, avoiding duplicates by title
    existing_titles = {t['title'] for t in tips}
    for fb in fallbacks:
        if len(tips) >= 5:
            break
        if fb['title'] not in existing_titles:
            tips.append(fb)
            existing_titles.add(fb['title'])

    return tips[:5]

# ─────────────────────────────────────────────────────────────────────────────
#  EXECUTIVE SUMMARY
# ─────────────────────────────────────────────────────────────────────────────
def render_exec_summary(f1, f2, fy1, fy2, client_name, advice):
    inc_chg  = f2['gross_income']  - f1['gross_income']
    tax_chg  = f2['display_gross_tax']  - f1['display_gross_tax']
    rate_chg = f2['display_eff_rate']    - f1['display_eff_rate']

    def arrow(v): return "▲" if v > 0 else ("▼" if v < 0 else "—")
    def col(v, good_if_negative=False):
        if v == 0: return ""
        pos = v > 0
        good = (not pos) if good_if_negative else pos
        return "delta-up" if good else "delta-down"

    _tips_filtered = [t for t in advice
                      if not (f2.get('super_cap_exceeded') and
                              'Maximise Concessional Super' in t.get('title',''))]
    high_tips = [t for t in _tips_filtered if t.get('priority') == 'High']
    tip_html  = "".join(f"<li>{t['title']}</li>" for t in high_tips[:3])

    st.markdown(f"""
    <div class="exec-box">
      <h2>📋 Executive Summary — {client_name}</h2>
      <div style="display:grid;grid-template-columns:repeat(auto-fit,minmax(200px,1fr));gap:16px;margin-bottom:16px;">
        <div><div style="font-size:11px;color:#888;text-transform:uppercase;">FY26 Gross Income</div>
             <div style="font-size:22px;font-weight:700;color:{DNA_BLUE};">${f2['gross_income']:,.0f}</div>
             <div class="{col(inc_chg)}">{arrow(inc_chg)} ${abs(inc_chg):,.0f} vs {fy1}</div></div>
        <div><div style="font-size:11px;color:#888;text-transform:uppercase;">FY26 Taxable Income{"&nbsp;&#10003;ATO" if f2.get("ato_verified") else ""}</div>
             <div style="font-size:22px;font-weight:700;color:{DNA_BLUE};">${f2['display_taxable']:,.0f}</div></div>
        <div><div style="font-size:11px;color:#888;text-transform:uppercase;">FY26 Tax Paid{"&nbsp;&#10003;ATO" if f2.get("ato_verified") else ""}</div>
             <div style="font-size:22px;font-weight:700;color:{DNA_RED};">${f2['display_gross_tax']:,.0f}</div>
             <div class="{col(tax_chg, good_if_negative=True)}">{arrow(tax_chg)} ${abs(tax_chg):,.0f} vs {fy1}</div></div>
        <div><div style="font-size:11px;color:#888;text-transform:uppercase;">Effective Rate</div>
             <div style="font-size:22px;font-weight:700;color:{DNA_BLUE};">{f2['display_eff_rate']}%</div>
             <div class="{col(rate_chg, good_if_negative=True)}">{arrow(rate_chg)} {abs(rate_chg):.1f}% vs {fy1}</div></div>
        <div><div style="font-size:11px;color:#888;text-transform:uppercase;">{"Refund" if f2["display_refund"]>0 else "Owing"}{"&nbsp;&#10003;ATO" if f2.get("ato_verified") else ""}</div>
             <div style="font-size:22px;font-weight:700;color:{''+DNA_GREEN if f2["display_refund"]>0 else DNA_RED};">
             {"${:,.0f}".format(f2['display_refund']) if f2['display_refund']>0 else "${:,.0f}".format(f2['display_owing'])}</div></div>
      </div>
      {"<div style='margin-top:10px;'><strong style='color:"+DNA_BLUE+";'>🎯 Top Priority Actions:</strong><ul style='margin:6px 0 0 16px;font-size:13px;color:#444;'>"+tip_html+"</ul></div>" if high_tips else ""}
    </div>
    """, unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
#  YOY COMPARISON TABLE
# ─────────────────────────────────────────────────────────────────────────────
def render_yoy_table(f1, f2, fy1, fy2):
    rows = [
        ("Gross Income",        f1['gross_income'],        f2['gross_income'],        False),
        ("Taxable Income",      f1['display_taxable'],     f2['display_taxable'],     False),
        ("Total Deductions",    f1['total_ded'],            f2['total_ded'],            True),
        ("Rental Net",          f1['rental_net'],           f2['rental_net'],           True),
        ("Income Tax",          f1['income_tax'],           f2['income_tax'],           False),
        ("Medicare Levy",       f1['medicare'],             f2['medicare'],             False),
        ("Offsets Applied",     f1['lito']+f1['seniors_off'],f2['lito']+f2['seniors_off'],True),
        ("Gross Tax",           f1['display_gross_tax'],    f2['display_gross_tax'],    False),
        ("PAYG Credits",        f1['credits'],              f2['credits'],              True),
        ("Refund / (Owing)",    f1['display_refund']-f1['display_owing'],  f2['display_refund']-f2['display_owing'],  True),
        ("Effective Rate (%)",  f1['display_eff_rate'],     f2['display_eff_rate'],     False),
    ]
    data = []
    for label, v1, v2, good_if_higher in rows:
        chg = v2 - v1
        if label == "Effective Rate (%)":
            r1 = f"{v1:.2f}%"; r2 = f"{v2:.2f}%"; rc = f"{chg:+.2f}%"
        else:
            r1 = f"${v1:,.0f}"; r2 = f"${v2:,.0f}"; rc = f"${chg:+,.0f}"
        trend = "▲" if chg > 0 else ("▼" if chg < 0 else "—")
        data.append({"Metric": label, fy1: r1, fy2: r2, "Change": rc, "Trend": trend})
    df = pd.DataFrame(data)
    st.dataframe(df, use_container_width=True, hide_index=True)

# ─────────────────────────────────────────────────────────────────────────────
#  PDF EXPORT
# ─────────────────────────────────────────────────────────────────────────────
def fig_bytes(fig, dpi=110):
    buf = io.BytesIO()
    fig.savefig(buf, format='PNG', dpi=dpi, bbox_inches='tight')
    buf.seek(0); return buf


def generate_pdf(client_name, fy1, fy2, f1, f2, advice, flags, figs: dict):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, HRFlowable, Image as RLImage,
                                    PageBreak, KeepTogether, BaseDocTemplate,
                                    PageTemplate, Frame)
    from reportlab.platypus.flowables import Flowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.utils import ImageReader
    import io as _io

    # ── Colours matching DNA template ─────────────────────────────────────────
    DBLUE  = colors.HexColor("#1B3A6B")
    DGOLD  = colors.HexColor("#C9A84C")
    DTEAL  = colors.HexColor("#006F73")
    DLTEAL = colors.HexColor("#6AABA5")
    DSKY   = colors.HexColor("#00B0F0")
    DGREEN = colors.HexColor("#27AE60")
    DRED   = colors.HexColor("#C0392B")
    LGREY  = colors.HexColor("#F0F4FA")
    WHITE  = colors.white

    PAGE_W, PAGE_H = A4  # 595 x 842 pts

    # ── Logo helper ────────────────────────────────────────────────────────────
    def _find_file(name):
        for p in [pathlib.Path(__file__).parent/name,
                  pathlib.Path(name),
                  pathlib.Path(__file__).resolve().parent/name]:
            if p.exists(): return str(p)
        return None

    LOGO_FILE     = _find_file("dna_logo.png")
    LOGO_TRANS    = _find_file("dna_logo_transparent.png")   # black→transparent version
    COVER_BG      = _find_file("dna_cover_bg.png")
    # Prefer transparent version for PDF (renders correctly on dark bg)
    LOGO_PDF      = LOGO_TRANS or LOGO_FILE
    HAS_LOGO      = bool(LOGO_PDF)
    HAS_COVER     = bool(COVER_BG)

    # ── Style helper ──────────────────────────────────────────────────────────
    def sty(name, **kw):
        return ParagraphStyle(name, **kw)

    # Shared styles
    h1s   = sty('h1s',  fontSize=13, textColor=DBLUE, fontName='Helvetica-Bold',
                        spaceBefore=16, spaceAfter=6, leading=18)
    h2s   = sty('h2s',  fontSize=11, textColor=DTEAL, fontName='Helvetica-Bold',
                        spaceBefore=10, spaceAfter=4, leading=15)
    bods  = sty('bods', fontSize=9.5, textColor=colors.HexColor("#333"),
                        leading=15, spaceAfter=5)
    bods_j= sty('bods_j',fontSize=9.5,textColor=colors.HexColor("#333"),
                        leading=15, spaceAfter=5, alignment=TA_JUSTIFY)
    subs  = sty('subs', fontSize=9, textColor=colors.HexColor("#555"),
                        leading=13, spaceAfter=3)
    cents = sty('cent', fontSize=9, textColor=colors.HexColor("#555"),
                        alignment=TA_CENTER, spaceAfter=2)
    advt  = sty('advt', fontSize=10.5, textColor=DBLUE, fontName='Helvetica-Bold',
                        spaceBefore=8, spaceAfter=3)
    advb  = sty('advb', fontSize=9.5, textColor=colors.HexColor("#444"),
                        leading=14, spaceAfter=2)
    savs  = sty('savs', fontSize=9,   textColor=DGREEN, fontName='Helvetica-Bold',
                        spaceAfter=6)
    riskt = sty('riskt',fontSize=9.5, textColor=colors.HexColor("#333"),
                        leading=14, spaceAfter=3)
    toc_s = sty('toc_s',fontSize=11,  textColor=DBLUE, fontName='Helvetica',
                        spaceBefore=6, spaceAfter=6, leading=16)
    disc  = sty('disc', fontSize=8.5, textColor=colors.HexColor("#777"),
                        leading=13, spaceAfter=4, alignment=TA_JUSTIFY)

    # ── Page canvas with header/footer ────────────────────────────────────────
    def make_canvas_cb(is_cover=False):
        def on_page(canv, doc):
            canv.saveState()
            if is_cover:
                # Single solid teal background — clean, no two-tone
                canv.setFillColor(colors.HexColor("#006F73"))
                canv.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
                # Subtle dark overlay on bottom third for depth
                canv.setFillColor(colors.HexColor("#00575A"))
                canv.rect(0, 0, PAGE_W, PAGE_H * 0.38, fill=1, stroke=0)
                # Gold accent bar at very bottom
                canv.setFillColor(DGOLD)
                canv.rect(0, 0, PAGE_W, 18, fill=1, stroke=0)
            else:
                # ── Running header (42pt tall) ──────────────────────────────
                HDR_H = 42
                canv.setFillColor(DBLUE)
                canv.rect(0, PAGE_H - HDR_H, PAGE_W, HDR_H, fill=1, stroke=0)
                # Logo — vertically centred in header, anchored left
                if HAS_LOGO:
                    try:
                        logo_h = 26
                        logo_w = 80   # wider to show full logo text
                        logo_y = PAGE_H - HDR_H/2 - logo_h/2
                        canv.drawImage(LOGO_PDF, 14*mm, logo_y,
                                       width=logo_w, height=logo_h,
                                       preserveAspectRatio=True, mask='auto')
                    except Exception:
                        canv.setFillColor(WHITE)
                        canv.setFont("Helvetica-Bold", 10)
                        canv.drawString(14*mm, PAGE_H - HDR_H/2 - 4, "DNA Accountants")
                else:
                    canv.setFillColor(WHITE)
                    canv.setFont("Helvetica-Bold", 10)
                    canv.drawString(14*mm, PAGE_H - HDR_H/2 - 4, "DNA Accountants Pty Ltd")

                # Header right — two lines vertically centred
                canv.setFillColor(WHITE)
                canv.setFont("Helvetica-Bold", 8.5)
                canv.drawRightString(PAGE_W - 14*mm, PAGE_H - HDR_H/2 + 2,
                    f"Tax Optimisation Report  |  {fy2}  |  {client_name}")
                canv.setFillColor(colors.HexColor("#BDD4F0"))
                canv.setFont("Helvetica", 7.5)
                canv.drawRightString(PAGE_W - 14*mm, PAGE_H - HDR_H/2 - 8,
                    "DNA Accountants Pty Ltd  |  Chartered Accountants")
                # Gold divider under header
                canv.setStrokeColor(DGOLD)
                canv.setLineWidth(2.5)
                canv.line(0, PAGE_H - HDR_H - 1, PAGE_W, PAGE_H - HDR_H - 1)

                # ── Running footer ──────────────────────────────────────────
                canv.setFillColor(DBLUE)
                canv.rect(0, 0, PAGE_W, 26, fill=1, stroke=0)
                canv.setFillColor(DGOLD)
                canv.rect(0, 25, PAGE_W, 2, fill=1, stroke=0)
                canv.setFillColor(WHITE)
                canv.setFont("Helvetica", 7.5)
                canv.drawString(14*mm, 9,
                    "DNA Accountants Pty Ltd  |  Chartered Accountants  |  110 Pitt Street, Sydney NSW 2000  |  "
                    "For guidance only — always seek professional tax advice.")
                canv.setFont("Helvetica-Bold", 8)
                canv.drawRightString(PAGE_W - 14*mm, 9, f"Page {doc.page}")
            canv.restoreState()
        return on_page

    # ── Document setup ────────────────────────────────────────────────────────
    buf = _io.BytesIO()

    # Cover page frame (no margins — full bleed)
    cover_frame = Frame(0, 0, PAGE_W, PAGE_H, leftPadding=0, rightPadding=0,
                        topPadding=0, bottomPadding=0)
    # Content page frame (under header, above footer)
    content_frame = Frame(14*mm, 30, PAGE_W - 28*mm, PAGE_H - 78,
                          leftPadding=0, rightPadding=0, topPadding=8, bottomPadding=0)

    cover_tpl   = PageTemplate(id='cover',   frames=[cover_frame],
                                onPage=make_canvas_cb(is_cover=True))
    content_tpl = PageTemplate(id='content', frames=[content_frame],
                                onPage=make_canvas_cb(is_cover=False))

    doc = BaseDocTemplate(buf, pagesize=A4, pageTemplates=[cover_tpl, content_tpl])

    story = []

    # ══════════════════════════════════════════════════════════════════════════
    #  PAGE 1 — COVER
    # ══════════════════════════════════════════════════════════════════════════
    from reportlab.platypus import NextPageTemplate

    # Cover content (positioned via absolute coords on canvas — use spacers to push down)
    story.append(NextPageTemplate('cover'))

    # Push content to lower half of cover page
    story.append(Spacer(1, PAGE_H * 0.30))

    def cover_para(text, font='Helvetica-Bold', size=26, color=WHITE, space_after=8,
                   alignment=TA_LEFT):
        return Paragraph(f'<font name="{font}" size="{size}">'
                         f'<font color="#{color.hexval()[2:] if hasattr(color,"hexval") else "FFFFFF"}">'
                         f'{text}</font></font>',
                         sty(f'cp_{size}', fontName=font, fontSize=size,
                             textColor=color, spaceAfter=space_after,
                             leftIndent=22*mm, leading=size*1.3, alignment=alignment))

    # Cover title block
    LM = 20*mm   # consistent left margin for all cover elements

    story.append(Paragraph(
        'DNA ACCOUNTANTS PTY LTD',
        sty('cvr_sub', fontName='Helvetica-Bold', fontSize=10, textColor=DGOLD,
            leftIndent=LM, spaceAfter=8)))
    story.append(Paragraph(
        'Tax Optimisation',
        sty('cvr_t1', fontName='Helvetica-Bold', fontSize=30, textColor=WHITE,
            leftIndent=LM, spaceAfter=0, leading=36)))
    story.append(Paragraph(
        'Report',
        sty('cvr_t2', fontName='Helvetica-Bold', fontSize=30, textColor=WHITE,
            leftIndent=LM, spaceAfter=14, leading=36)))

    # Gold rule — full width
    story.append(HRFlowable(width="100%", thickness=3, color=DGOLD, spaceAfter=28))

    story.append(Paragraph(
        f'Prepared for: <b>{client_name}</b>',
        sty('cvr_cl', fontName='Helvetica', fontSize=13, textColor=colors.HexColor("#E8F4F0"),
            leftIndent=LM, spaceAfter=6)))
    story.append(Paragraph(
        f'{fy1} &amp; {fy2} Financial Years',
        sty('cvr_fy', fontName='Helvetica', fontSize=11, textColor=colors.HexColor("#BDD4F0"),
            leftIndent=LM, spaceAfter=5)))
    story.append(Paragraph(
        f'Date: {datetime.now().strftime("%d %B %Y")}',
        sty('cvr_dt', fontName='Helvetica', fontSize=10, textColor=colors.HexColor("#9ABFD8"),
            leftIndent=LM, spaceAfter=12)))

    # AI + reviewed badges
    badge_data = [
        [Paragraph('<font name="Helvetica-Bold" size="9" color="#FFFFFF">⚙ Drafted by AI</font>',
                   sty('b1', fontName='Helvetica-Bold', fontSize=9, textColor=WHITE,
                       alignment=TA_CENTER)),
         Paragraph('<font name="Helvetica-Bold" size="9" color="#FFFFFF">✓ Reviewed &amp; Signed Off by DNA Team</font>',
                   sty('b2', fontName='Helvetica-Bold', fontSize=9, textColor=WHITE,
                       alignment=TA_CENTER))]
    ]
    badge_tbl = Table(badge_data, colWidths=[5*cm, 8*cm])
    badge_tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0,0),(0,0), colors.HexColor("#006F73")),
        ("BACKGROUND",   (1,0),(1,0), DBLUE),
        ("TOPPADDING",   (0,0),(-1,-1), 8),
        ("BOTTOMPADDING",(0,0),(-1,-1), 8),
        ("LEFTPADDING",  (0,0),(-1,-1), 12),
        ("RIGHTPADDING", (0,0),(-1,-1), 12),
    ]))
    story.append(Spacer(1, 8))
    # Indent badge table to match cover left margin
    badge_wrapper = Table([['  ', badge_tbl]], colWidths=[20*mm, 13*cm])
    badge_wrapper.setStyle(TableStyle([
        ("LEFTPADDING",  (0,0),(-1,-1), 0),
        ("RIGHTPADDING", (0,0),(-1,-1), 0),
        ("TOPPADDING",   (0,0),(-1,-1), 0),
        ("BOTTOMPADDING",(0,0),(-1,-1), 0),
        ("VALIGN",       (0,0),(-1,-1), "MIDDLE"),
    ]))
    story.append(KeepTogether([badge_wrapper]))

    # Logo at bottom of cover
    if HAS_LOGO:
        story.append(Spacer(1, 20))
        try:
            story.append(RLImage(LOGO_PDF, width=5*cm, height=2*cm, kind='proportional'))
        except Exception:
            pass

    # ══════════════════════════════════════════════════════════════════════════
    #  PAGE 2 — TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════════
    story.append(NextPageTemplate('content'))
    story.append(PageBreak())

    story.append(Paragraph("Table of Contents", h1s))
    story.append(HRFlowable(width="100%", thickness=2.5, color=DGOLD, spaceAfter=6))
    story.append(Paragraph("DNA Accountants Pty Ltd",
                            sty('toc_co', fontName='Helvetica', fontSize=10,
                                textColor=DTEAL, spaceAfter=12)))

    toc_items = [
        ("1.", "Basis of Preparation",               "3"),
        ("2.", "Executive Summary",                   "4"),
        ("3.", "Year-on-Year Financial Comparison",   "5"),
        ("4.", "Income & Expense Analysis",           "6"),
        ("5.", "Tax Bracket Position",               "7"),
        ("6.", "Rental Property Analysis",           "8"),
        ("7.", "Super Contribution Modeller",        "9"),
        ("8.", "CGT Harvest Opportunity",            "10"),
        ("9.", "ATO Audit Risk Indicators",          "11"),
        ("10.", "AI Tax Optimisation Advice",        "12"),
        ("11.", "Disclaimer",                        "13"),
    ]
    for num, title, pg in toc_items:
        # Dotted leader line
        toc_row = Table(
            [[Paragraph(f'<b>{num}</b>', sty(f'tn{num}', fontName='Helvetica-Bold',
                        fontSize=11, textColor=DBLUE)),
              Paragraph(title, toc_s),
              Paragraph(pg, sty(f'tp{num}', fontName='Helvetica', fontSize=11,
                                textColor=DTEAL, alignment=TA_RIGHT))]],
            colWidths=[1.2*cm, 13*cm, 1.8*cm]
        )
        toc_row.setStyle(TableStyle([
            ("VALIGN",       (0,0),(-1,-1), "MIDDLE"),
            ("BOTTOMPADDING",(0,0),(-1,-1), 8),
            ("TOPPADDING",   (0,0),(-1,-1), 2),
            ("LINEBELOW",    (0,0),(-1,0),  0.5, colors.HexColor("#DDDDDD")),
            ("LEFTPADDING",  (0,0),(-1,-1), 0),
            ("RIGHTPADDING", (0,0),(-1,-1), 0),
        ]))
        story.append(toc_row)

    # ══════════════════════════════════════════════════════════════════════════
    #  PAGE 3 — BASIS OF PREPARATION
    # ══════════════════════════════════════════════════════════════════════════
    story.append(PageBreak())
    story.append(Paragraph("1. Basis of Preparation", h1s))
    story.append(HRFlowable(width="100%", thickness=2.5, color=DGOLD, spaceAfter=4))
    story.append(Paragraph("DNA Accountants Pty Ltd",
                            sty('sec_co1', fontName='Helvetica', fontSize=10,
                                textColor=DTEAL, spaceAfter=10)))

    story.append(Paragraph("About This Report", h2s))
    story.append(Paragraph(
        f"This Tax Optimisation Report has been prepared by <b>DNA Accountants Pty Ltd</b> "
        f"(Chartered Accountants, Sydney NSW) for <b>{client_name}</b> in relation to the "
        f"{fy1} and {fy2} financial years. The report has been generated using DNA's "
        f"AI-assisted tax analysis platform and reviewed and signed off by a registered tax "
        f"agent in accordance with the firm's quality management procedures.",
        bods_j))

    story.append(Spacer(1, 8))
    story.append(Paragraph("AI Use — Alignment with TPB(I) D62/2026", h2s))

    # TPB alignment notice box
    tpb_box = Table([[Paragraph(
        "<b>Tax Practitioners Board (TPB) Exposure Draft TPB(I) D62/2026</b><br/><br/>"
        "This report has been prepared in alignment with the TPB's draft Information Sheet "
        "TPB(I) D62/2026 (<i>The use of Artificial Intelligence and the Code of Professional "
        "Conduct</i>, issued 23 March 2026). The key principles applied are:<br/><br/>"
        "<b>1. Competency (Code Items 7–10 &amp; Determination ss.30, 35, 40):</b> "
        "AI-generated analysis has been used to assist — not replace — the professional "
        "judgement of a registered tax agent. All AI outputs have been reviewed, verified "
        "and supplemented with expert analysis before inclusion in this report.<br/><br/>"
        "<b>2. Confidentiality (Code Item 6):</b> Client information has been processed "
        "within DNA's secure, privacy-compliant AI environment. No client data has been "
        "disclosed to third-party AI systems without appropriate consent and data governance "
        "controls in place. All processing complies with the <i>Privacy Act 1988</i> (Cth) "
        "and the Australian Privacy Principles (APPs).<br/><br/>"
        "<b>3. Accuracy &amp; Reasonable Care (Code Items 9–10):</b> AI outputs have been "
        "independently verified against current ATO legislation and guidelines. The tax agent "
        "has applied professional judgement to assess, contest and validate all AI "
        "recommendations before delivery to the client.<br/><br/>"
        "<b>4. Quality Management (Determination s.40):</b> DNA maintains a documented "
        "system of quality management that governs AI use, including review protocols, "
        "data security standards and staff competency requirements for AI-assisted services.",
        sty('tpb_body', fontSize=9, textColor=colors.HexColor("#1A1A1A"),
            leading=14, spaceAfter=0))]],
        colWidths=[PAGE_W - 28*mm - 12])
    tpb_box.setStyle(TableStyle([
        ("BACKGROUND",   (0,0),(-1,-1), colors.HexColor("#EAF5F5")),
        ("LEFTPADDING",  (0,0),(-1,-1), 14),
        ("RIGHTPADDING", (0,0),(-1,-1), 14),
        ("TOPPADDING",   (0,0),(-1,-1), 12),
        ("BOTTOMPADDING",(0,0),(-1,-1), 12),
        ("LINEBEFORE",   (0,0),(0,-1),  5, DTEAL),
    ]))
    story.append(tpb_box)
    story.append(Spacer(1, 10))

    story.append(Paragraph("Scope &amp; Limitations", h2s))
    story.append(Paragraph(
        "This report is prepared on the basis of financial information provided by the client "
        "for the relevant financial years. The analysis and recommendations are intended as "
        "general guidance and do not constitute legal or financial advice. Tax laws and ATO "
        "interpretations are subject to change. Clients should discuss all recommendations "
        "with their DNA advisor before taking any action. DNA Accountants Pty Ltd accepts no "
        "liability for decisions made solely on the basis of this report.",
        bods_j))

    story.append(Spacer(1, 8))
    story.append(Paragraph("Data Currency", h2s))
    prep_data = [
        ["Report prepared:", datetime.now().strftime("%d %B %Y")],
        ["Financial years covered:", f"{fy1} and {fy2}"],
        ["Tax brackets applied:", "FY2025-26 ATO rates (effective 1 July 2025)"],
        ["Concessional super cap:", "$30,000 (FY2025-26)"],
        ["Medicare levy rate:", "2.0%"],
        ["AI system:", "Google Gemini Flash 2.5 (reviewed by DNA team)"],
        ["TPB alignment:", "TPB(I) D62/2026 — AI & Code of Professional Conduct"],
    ]
    prep_tbl = Table(prep_data, colWidths=[5*cm, PAGE_W - 28*mm - 12 - 5*cm])
    prep_tbl.setStyle(TableStyle([
        ("FONTNAME",     (0,0),(0,-1), "Helvetica-Bold"),
        ("FONTSIZE",     (0,0),(-1,-1), 9),
        ("TEXTCOLOR",    (0,0),(0,-1), DBLUE),
        ("BACKGROUND",   (0,0),(-1,0), colors.HexColor("#EAF0FB")),
        ("BACKGROUND",   (0,2),(-1,2), colors.HexColor("#EAF0FB")),
        ("BACKGROUND",   (0,4),(-1,4), colors.HexColor("#EAF0FB")),
        ("BACKGROUND",   (0,6),(-1,6), colors.HexColor("#EAF0FB")),
        ("GRID",         (0,0),(-1,-1), 0.4, colors.HexColor("#CCCCCC")),
        ("TOPPADDING",   (0,0),(-1,-1), 5),
        ("BOTTOMPADDING",(0,0),(-1,-1), 5),
        ("LEFTPADDING",  (0,0),(-1,-1), 8),
        ("RIGHTPADDING", (0,0),(-1,-1), 8),
    ]))
    story.append(prep_tbl)

    # ══════════════════════════════════════════════════════════════════════════
    #  PAGE 4 — EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    story.append(PageBreak())
    story.append(Paragraph("2. Executive Summary", h1s))
    story.append(HRFlowable(width="100%", thickness=2.5, color=DGOLD, spaceAfter=4))
    story.append(Paragraph('DNA Accountants Pty Ltd',
                            sty('sec_co2', fontName='Helvetica', fontSize=10,
                                textColor=DTEAL, spaceAfter=10)))

    inc_chg = f2['gross_income'] - f1['gross_income']
    tax_chg = f2['display_gross_tax'] - f1['display_gross_tax']
    es_data = [
        ["Metric", fy1, fy2, "Change"],
        ["Gross Income",
         f"${f1['gross_income']:,.0f}", f"${f2['gross_income']:,.0f}",
         "${:+,.0f}".format(inc_chg)],
        ["Taxable Income",
         f"${f1['display_taxable']:,.0f}", f"${f2['display_taxable']:,.0f}",
         "${:+,.0f}".format(f2['display_taxable']-f1['display_taxable'])],
        ["Gross Tax",
         f"${f1['display_gross_tax']:,.0f}", f"${f2['display_gross_tax']:,.0f}",
         "${:+,.0f}".format(tax_chg)],
        ["Effective Rate",
         f"{f1['display_eff_rate']}%", f"{f2['display_eff_rate']}%",
         "{:+.1f}%".format(f2['display_eff_rate']-f1['display_eff_rate'])],
        ["Refund / (Owing)",
         ("${:,.0f}".format(f1['display_refund']) if f1['display_refund']>0 else "(${:,.0f})".format(f1['display_owing'])),
         ("${:,.0f}".format(f2['display_refund']) if f2['display_refund']>0 else "(${:,.0f})".format(f2['display_owing'])),
         "\u2713 ATO verified" if f2.get("ato_verified") else ""],
    ]
    # ── ATO Tax Breakdown (shown only when verified figures provided) ──
    if f2.get("ato_verified"):
        es_data += [
            ["" , "", "", ""],
            ["Income Tax (ATO)",
             f"${f1['display_income_tax']:,.2f}", f"${f2['display_income_tax']:,.2f}", ""],
            ["Medicare Levy (ATO)",
             f"${f1['display_medicare']:,.2f}", f"${f2['display_medicare']:,.2f}", ""],
            ["Offsets / Credits (ATO)",
             f"(${f1['ato_other_offsets']:,.2f})", f"(${f2['ato_other_offsets']:,.2f})", ""],
            ["Gross Tax (ATO)",
             f"${f1['display_gross_tax']:,.2f}", f"${f2['display_gross_tax']:,.2f}", ""],
            ["PAYG Credits",
             f"(${f1['credits']:,.2f})", f"(${f2['credits']:,.2f})", ""],
        ]
    es_data += [
    ]
    et = Table(es_data, colWidths=[5.5*cm, 3.2*cm, 3.2*cm, 3.2*cm])
    et.setStyle(TableStyle([
        ("BACKGROUND",   (0,0),(-1,0), DBLUE),
        ("TEXTCOLOR",    (0,0),(-1,0), WHITE),
        ("FONTNAME",     (0,0),(-1,0), "Helvetica-Bold"),
        ("FONTNAME",     (0,1),(0,-1), "Helvetica-Bold"),
        ("TEXTCOLOR",    (0,1),(0,-1), DBLUE),
        ("FONTSIZE",     (0,0),(-1,-1), 9.5),
        ("BACKGROUND",   (0,1),(-1,1), WHITE),
        ("BACKGROUND",   (0,2),(-1,2), LGREY),
        ("BACKGROUND",   (0,3),(-1,3), WHITE),
        ("BACKGROUND",   (0,4),(-1,4), LGREY),
        ("BACKGROUND",   (0,5),(-1,5), colors.HexColor("#EAF5F5")),
        ("GRID",         (0,0),(-1,-1), 0.4, colors.HexColor("#CCCCCC")),
        ("ALIGN",        (1,0),(-1,-1), "RIGHT"),
        ("TOPPADDING",   (0,0),(-1,-1), 6),
        ("BOTTOMPADDING",(0,0),(-1,-1), 6),
        ("LEFTPADDING",  (0,0),(-1,-1), 8),
        ("RIGHTPADDING", (0,0),(-1,-1), 8),
    ]))
    story += [et, Spacer(1, 14)]

    # High priority tips
    high_tips = [t for t in advice if t.get('priority')=='High']
    if high_tips:
        story.append(Paragraph("Top Priority Tax Actions", h2s))
        for tip in high_tips[:3]:
            story.append(Paragraph(
                f"<b>▶ {tip['title']}</b> — {tip.get('detail','')[:200]}...",
                sty('htip', fontSize=9.5, textColor=colors.HexColor("#333"),
                    leading=14, spaceAfter=6,
                    leftIndent=10, borderPad=6)))

    # ══════════════════════════════════════════════════════════════════════════
    #  CHARTS PAGES
    # ══════════════════════════════════════════════════════════════════════════
    def add_chart_page(section_num, title, fig_key, caption=""):
        story.append(PageBreak())
        story.append(Paragraph(f"{section_num}. {title}", h1s))
        story.append(HRFlowable(width="100%", thickness=2.5, color=DGOLD, spaceAfter=4))
        story.append(Paragraph("DNA Accountants Pty Ltd",
                                sty(f'sec_coc{section_num}', fontName='Helvetica', fontSize=10,
                                    textColor=DTEAL, spaceAfter=8)))
        fig = figs.get(fig_key)
        if fig:
            story.append(RLImage(fig_bytes(fig), width=15.5*cm, height=8.2*cm))
        if caption:
            story.append(Paragraph(caption, subs))
        story.append(Spacer(1, 8))

    add_chart_page("3", "Year-on-Year Financial Comparison", 'yoy',
        f"Comparison of key financial metrics between {fy1} and {fy2}.")

    # Income & Expense on same page (two charts)
    story.append(PageBreak())
    story.append(Paragraph("4. Income & Expense Analysis", h1s))
    story.append(HRFlowable(width="100%", thickness=2.5, color=DGOLD, spaceAfter=4))
    story.append(Paragraph("DNA Accountants Pty Ltd",
                            sty('sec_co4', fontName='Helvetica', fontSize=10,
                                textColor=DTEAL, spaceAfter=8)))
    for fk in ['income','expense']:
        fig = figs.get(fk)
        if fig:
            story.append(RLImage(fig_bytes(fig), width=15.5*cm, height=6.5*cm))
            story.append(Spacer(1, 6))

    add_chart_page("5", "Tax Bracket Position", 'bracket',
        f"Client taxable income position relative to ATO tax brackets — {fy2}.")

    if figs.get('rental'):
        add_chart_page("6", "Rental Property Cash Flow Analysis", 'rental',
            "Rental income, expenses, interest, depreciation, net rental and estimated tax benefit.")
    else:
        story.append(PageBreak())
        story.append(Paragraph("6. Rental Property Analysis", h1s))
        story.append(HRFlowable(width="100%", thickness=2.5, color=DGOLD, spaceAfter=10))
        story.append(Paragraph("No rental income declared for this client.", bods))

    if figs.get('super'):
        add_chart_page("7", "Super Contribution Tax Saving Modeller", 'super',
            f"Estimated tax saving at each additional super contribution level. "
            f"Marginal rate {int(f2['marginal_rate']*100)}% vs 15% super tax = "
            f"{int((f2['marginal_rate']-0.15)*100)}c saving per $1 contributed.")
    else:
        story.append(PageBreak())
        story.append(Paragraph("7. Super Contribution Modeller", h1s))
        story.append(HRFlowable(width="100%", thickness=2.5, color=DGOLD, spaceAfter=10))
        story.append(Paragraph("Concessional super cap is fully utilised for this year.", bods))

    if figs.get('cgt'):
        add_chart_page("8", "CGT Harvest Opportunity", 'cgt',
            "Capital gains, available losses, unrealised harvestable losses and estimated tax impact.")
    else:
        story.append(PageBreak())
        story.append(Paragraph("8. CGT Harvest Opportunity", h1s))
        story.append(HRFlowable(width="100%", thickness=2.5, color=DGOLD, spaceAfter=10))
        story.append(Paragraph("No capital gains or unrealised losses declared.", bods))

    # ══════════════════════════════════════════════════════════════════════════
    #  AUDIT FLAGS
    # ══════════════════════════════════════════════════════════════════════════
    story.append(PageBreak())
    story.append(Paragraph("9. ATO Audit Risk Indicators", h1s))
    story.append(HRFlowable(width="100%", thickness=2.5, color=DGOLD, spaceAfter=8))
    story.append(Paragraph(
        "The following indicators are based on ATO benchmarks and statistical analysis. "
        "These are not predictions of an audit but highlight areas requiring attention.",
        bods_j))
    story.append(Spacer(1, 8))

    col_map    = {"High":"#FFF0F0","Medium":"#FFFBF0","Low":"#F0FFF4"}
    border_map = {"High":DRED,     "Medium":DGOLD,    "Low":DGREEN}
    for level, msg in flags:
        flag_tbl = Table(
            [[Paragraph(f"<b>{level} Risk:</b>  {msg}", riskt)]],
            colWidths=[PAGE_W - 28*mm - 12])
        flag_tbl.setStyle(TableStyle([
            ("BACKGROUND",   (0,0),(-1,-1), colors.HexColor(col_map.get(level,'#F5F5F5'))),
            ("LEFTPADDING",  (0,0),(-1,-1), 10),
            ("RIGHTPADDING", (0,0),(-1,-1), 10),
            ("TOPPADDING",   (0,0),(-1,-1), 7),
            ("BOTTOMPADDING",(0,0),(-1,-1), 7),
            ("LINEBEFORE",   (0,0),(0,-1),  5, border_map.get(level, DGOLD)),
        ]))
        story += [flag_tbl, Spacer(1, 6)]

    # ══════════════════════════════════════════════════════════════════════════
    #  AI ADVICE
    # ══════════════════════════════════════════════════════════════════════════
    story.append(PageBreak())
    story.append(Paragraph("10. AI Tax Optimisation Advice", h1s))
    story.append(HRFlowable(width="100%", thickness=2.5, color=DGOLD, spaceAfter=8))
    story.append(Paragraph(
        "The following five recommendations are generated through AI analysis of your "
        f"two-year financial profile ({fy1} &amp; {fy2}) and reviewed by a DNA tax advisor. "
        "They are intended as a starting point for discussion — please consult your DNA "
        "advisor before acting on any suggestion.",
        bods_j))
    story.append(Spacer(1, 10))

    pri_col = {"High": DRED, "Medium": DGOLD, "Low": DGREEN}
    for i, tip in enumerate(advice, 1):
        pri = tip.get('priority','Medium')
        pc  = pri_col.get(pri, DGOLD)
        tip_tbl = Table([
            [Paragraph(f"<b>{i}</b>",
                       sty(f'anum{i}', fontName='Helvetica-Bold', fontSize=16,
                           textColor=pc, alignment=TA_CENTER)),
             [Paragraph(f"<b>{tip.get('title','')}</b>",
                        sty(f'atit{i}', fontName='Helvetica-Bold', fontSize=11,
                            textColor=DBLUE, spaceAfter=4)),
              Paragraph(tip.get('detail',''),
                        sty(f'adet{i}', fontSize=9.5, textColor=colors.HexColor("#333"),
                            leading=14, spaceAfter=4)),
              Paragraph(
                  f"<font color='#27AE60'><b>Estimated Saving: "
                  f"{tip.get('estimated_saving','N/A')}</b></font>"
                  f"  &nbsp;&nbsp;  "
                  f"<font color='#{pc.hexval()[2:]}'><b>Priority: {pri}</b></font>",
                  sty(f'asav{i}', fontSize=9, leading=12))]]
        ], colWidths=[1.2*cm, PAGE_W - 28*mm - 12 - 1.2*cm])
        tip_tbl.setStyle(TableStyle([
            ("BACKGROUND",   (0,0),(-1,-1), WHITE),
            ("LEFTPADDING",  (0,0),(0,-1),  0),
            ("RIGHTPADDING", (0,0),(-1,-1), 8),
            ("TOPPADDING",   (0,0),(-1,-1), 10),
            ("BOTTOMPADDING",(0,0),(-1,-1), 10),
            ("VALIGN",       (0,0),(0,-1),  "MIDDLE"),
            ("LINEBEFORE",   (0,0),(0,-1),  5, pc),
            ("LINEBELOW",    (0,0),(-1,-1), 0.5, colors.HexColor("#DDDDDD")),
        ]))
        story.append(tip_tbl)
        story.append(Spacer(1, 4))

    # ══════════════════════════════════════════════════════════════════════════
    #  DISCLAIMER
    # ══════════════════════════════════════════════════════════════════════════
    story.append(PageBreak())
    story.append(Paragraph("11. Disclaimer", h1s))
    story.append(HRFlowable(width="100%", thickness=2.5, color=DGOLD, spaceAfter=4))
    story.append(Paragraph('DNA Accountants Pty Ltd',
                            sty('sec_co11', fontName='Helvetica', fontSize=10,
                                textColor=DTEAL, spaceAfter=10)))

    disc_items = [
        ("General Advice Warning",
         "This report contains general tax information and is not personal financial or legal "
         "advice. The information is based on the data provided and current Australian tax law "
         "as at the date of preparation. DNA Accountants Pty Ltd recommends you seek "
         "independent professional advice before making any financial or tax decisions."),
        ("AI-Generated Content",
         "Portions of this report have been generated using artificial intelligence tools, "
         "in accordance with TPB(I) D62/2026. All AI-generated content has been reviewed "
         "and verified by a registered tax agent at DNA Accountants Pty Ltd prior to "
         "delivery. The use of AI does not diminish our professional obligations or the "
         "quality of advice provided."),
        ("Accuracy of Information",
         "The analysis in this report is based solely on information provided to DNA "
         "Accountants Pty Ltd. If the information provided is incomplete, inaccurate or "
         "has changed, the recommendations may not be applicable. Clients are responsible "
         "for ensuring all information provided is current and complete."),
        ("Privacy",
         "Client information has been handled in accordance with the Privacy Act 1988 (Cth) "
         "and DNA Accountants' Privacy Policy. Client data used in AI processing has been "
         "handled within a secure, privacy-compliant environment with appropriate consent "
         "and data governance controls."),
        ("Limitation of Liability",
         "DNA Accountants Pty Ltd accepts no liability for any loss or damage arising "
         "from reliance on this report without first obtaining professional advice tailored "
         "to your specific circumstances."),
    ]
    for title, text in disc_items:
        story.append(Paragraph(title, h2s))
        story.append(Paragraph(text, disc))
        story.append(Spacer(1, 4))

    story.append(Spacer(1, 16))
    story.append(HRFlowable(width="100%", thickness=2, color=DGOLD))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        f"<b>DNA Accountants Pty Ltd</b>  |  Chartered Accountants  |  110 Pitt Street, Sydney NSW 2000<br/>"
        f"Registered Tax Agent  |  110 Pitt Street, Sydney NSW 2000  |  (02) 9064 4400<br/>"
        f"<i>Prepared: {datetime.now().strftime('%d %B %Y')} — This document is confidential "
        f"and intended solely for {client_name}.</i>",
        sty('fin', fontSize=8.5, textColor=colors.HexColor("#444"),
            leading=13, alignment=TA_CENTER)))

    doc.build(story)
    buf.seek(0)
    return buf.read()

# ─────────────────────────────────────────────────────────────────────────────
#  WORD EXPORT

# ─────────────────────────────────────────────────────────────────────────────
def generate_docx(client_name, fy1, fy2, f1, f2, advice, flags, figs):
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLUE  = RGBColor(0x1B,0x3A,0x6B); GOLD  = RGBColor(0xC9,0xA8,0x4C)
    WHITE = RGBColor(0xFF,0xFF,0xFF); GREEN = RGBColor(0x27,0xAE,0x60)
    GREY  = RGBColor(0x55,0x55,0x55); RED   = RGBColor(0xC0,0x39,0x2B)

    def bg(cell, hex_):
        tc=cell._tc; pr=tc.get_or_add_tcPr()
        s=OxmlElement('w:shd')
        s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hex_)
        pr.append(s)

    def run(para, text, bold=False, italic=False, size=10, color=None):
        r=para.add_run(text); r.bold=bold; r.italic=italic
        r.font.name='Arial'; r.font.size=Pt(size)
        if color: r.font.color.rgb=color
        return r

    def sec(doc, text, size=13, col=None, sb=12, sa=5):
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
        run(p, text, bold=True, size=size, color=col or BLUE)
        pr=p._p.get_or_add_pPr(); bd=OxmlElement('w:pBdr')
        b=OxmlElement('w:bottom')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
        b.set(qn('w:space'),'1'); b.set(qn('w:color'),'C9A84C')
        bd.append(b); pr.append(bd)

    def ins_fig(doc, fig, w_cm=15):
        if fig is None: return
        buf=io.BytesIO(); fig.savefig(buf,format='PNG',dpi=100,bbox_inches='tight'); buf.seek(0)
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after=Pt(8)
        p.add_run().add_picture(buf, width=Cm(w_cm))

    doc = DocxDoc()
    for section in doc.sections:
        section.top_margin=Cm(2); section.bottom_margin=Cm(2)
        section.left_margin=Cm(2.5); section.right_margin=Cm(2.5)

    # Banner
    LOGO_FILE = str(pathlib.Path(__file__).parent / "dna_logo.png")
    ht = doc.add_table(rows=1, cols=2); ht.style='Table Grid'
    lc = ht.rows[0].cells[0]; lc.width=Cm(4.8)
    bg(lc, "1B3A6B"); lp=lc.paragraphs[0]
    lp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before=Pt(6); lp.paragraph_format.space_after=Pt(6)
    if os.path.exists(LOGO_FILE):
        lp.add_run().add_picture(LOGO_FILE, width=Cm(4.0))
    else:
        run(lp,"DNA",bold=True,size=28,color=WHITE)

    rc = ht.rows[0].cells[1]; rc.width=Cm(10.8)
    bg(rc,"1B3A6B"); rc.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    rc.paragraphs[0]._p.getparent().remove(rc.paragraphs[0]._p)

    def cp(cell, text, bold=False, size=11, color=WHITE, sb=0, sa=4):
        p=OxmlElement('w:p'); cell._tc.append(p)
        from docx.text.paragraph import Paragraph as DP
        para=DP(p,cell); para.paragraph_format.space_before=Pt(sb); para.paragraph_format.space_after=Pt(sa)
        run(para,text,bold=bold,size=size,color=color); return para

    cp(rc,"DNA Accountants Pty Ltd",      bold=True,size=17,sb=8,sa=5)
    cp(rc,"Chartered Accountants  |  Sydney, Australia",size=10,color=RGBColor(0xBD,0xD4,0xF0),sa=4)
    cp(rc,f"Tax Optimisation Report  —  {fy1} & {fy2}",size=10,color=RGBColor(0xD0,0xE2,0xF5),sa=9)

    # Gold divider
    gd=doc.add_paragraph(); gd.paragraph_format.space_before=Pt(0); gd.paragraph_format.space_after=Pt(0)
    pr=gd._p.get_or_add_pPr(); bd=OxmlElement('w:pBdr')
    b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'12')
    b.set(qn('w:space'),'1'); b.set(qn('w:color'),'C9A84C')
    bd.append(b); pr.append(bd)

    cl=doc.add_paragraph(); cl.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cl.paragraph_format.space_before=Pt(8); cl.paragraph_format.space_after=Pt(14)
    run(cl,"Prepared for: ",size=11,color=GREY); run(cl,client_name,bold=True,size=11,color=BLUE)
    run(cl,f"     |     Date: {datetime.now().strftime('%d %B %Y')}",size=11,color=GREY)

    # Exec summary table
    sec(doc,"Executive Summary")
    rows=[("Gross Income",f1['gross_income'],f2['gross_income']),
          ("Taxable Income",f1['taxable_income'],f2['taxable_income']),
          ("Gross Tax",f1['gross_tax'],f2['gross_tax']),
          ("Effective Rate",f1['eff_rate'],f2['eff_rate']),
          ("Refund/(Owing)",f1['refund']-f1['owing'],f2['refund']-f2['owing'])]
    t=doc.add_table(rows=1+len(rows),cols=5); t.style='Table Grid'
    for j,(h,cw) in enumerate(zip(["Metric",fy1,fy2,"Change","Trend"],
                                   [Cm(4.5),Cm(3.2),Cm(3.2),Cm(3.2),Cm(1.5)]),0):
        c=t.rows[0].cells[j]; bg(c,"1B3A6B"); c.width=cw
        run(c.paragraphs[0],h,bold=True,size=10,color=WHITE)
    for i,(label,v1,v2) in enumerate(rows,1):
        chg=v2-v1; tr="▲" if chg>0 else ("▼" if chg<0 else "—")
        is_rate=label=="Effective Rate"
        r1=f"{v1:.1f}%" if is_rate else f"${v1:,.0f}"
        r2=f"{v2:.1f}%" if is_rate else f"${v2:,.0f}"
        rc2=f"{chg:+.1f}%" if is_rate else f"${chg:+,.0f}"
        row=t.rows[i]
        bg(row.cells[0],"D5E8F0")
        run(row.cells[0].paragraphs[0],label,bold=True,size=10,color=BLUE)
        for j,txt in enumerate([r1,r2,rc2,tr],1):
            p=row.cells[j].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            run(p,txt,size=10)
    doc.add_paragraph()

    # Charts
    if 'yoy' in figs: sec(doc,"Year-on-Year Comparison"); ins_fig(doc, figs.get('yoy'))
    if 'income' in figs: sec(doc,f"Income Breakdown — {fy2}"); ins_fig(doc, figs.get('income'))
    if 'expense' in figs: sec(doc,f"Expenses & Deductions — {fy2}"); ins_fig(doc, figs.get('expense'))
    if 'bracket' in figs: sec(doc,"Tax Bracket Position"); ins_fig(doc, figs.get('bracket'))
    if 'rental' in figs and figs['rental']: sec(doc,"Rental Cash Flow"); ins_fig(doc, figs.get('rental'))
    if 'super' in figs and figs['super']: sec(doc,"Super Contribution Modeller"); ins_fig(doc, figs.get('super'))
    if 'cgt' in figs and figs['cgt']: sec(doc,"CGT Harvest Analysis"); ins_fig(doc, figs.get('cgt'))

    # Audit flags
    sec(doc,"ATO Audit Risk Indicators")
    for level, msg in flags:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); p.paragraph_format.left_indent=Cm(0.5)
        col_map={"High":RED,"Medium":GOLD,"Low":GREEN}
        run(p,f"{level} Risk: ",bold=True,size=10,color=col_map.get(level,GOLD))
        run(p,msg,size=10)

    # AI advice
    sec(doc,"AI-Powered Tax Optimisation Advice")
    disc=doc.add_paragraph()
    run(disc,"Based on two-year trend analysis. Discuss with your DNA advisor before acting.",
        italic=True,size=9.5,color=GREY)
    disc.paragraph_format.space_after=Pt(8)

    for i,tip in enumerate(advice,1):
        pri=tip.get('priority','Medium')
        pri_col={"High":RED,"Medium":GOLD,"Low":GREEN}.get(pri,GOLD)
        tp=doc.add_paragraph(); tp.paragraph_format.space_before=Pt(8); tp.paragraph_format.space_after=Pt(2)
        run(tp,f"{i}.  {tip.get('title','')}",bold=True,size=12,color=BLUE)
        dp=doc.add_paragraph(); dp.paragraph_format.space_after=Pt(2); dp.paragraph_format.left_indent=Cm(0.6)
        run(dp,tip.get('detail',''),size=10)
        sp=doc.add_paragraph(); sp.paragraph_format.space_after=Pt(6); sp.paragraph_format.left_indent=Cm(0.6)
        run(sp,f"Estimated Saving: {tip.get('estimated_saving','N/A')}  |  Priority: {pri}",
            bold=True,size=10,color=GREEN)

    # Footer
    doc.add_paragraph()
    fp=doc.add_paragraph(); fp.paragraph_format.space_before=Pt(12); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pr=fp._p.get_or_add_pPr(); bd=OxmlElement('w:pBdr')
    b=OxmlElement('w:top'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'1'); b.set(qn('w:color'),'C9A84C'); bd.append(b); pr.append(bd)
    run(fp,"DNA Accountants Pty Ltd | Chartered Accountants | 110 Pitt Street, Sydney NSW 2000 | "
        "This report is for general guidance only. Always seek professional tax advice.",
        italic=True,size=8,color=GREY)

    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()

# ─────────────────────────────────────────────────────────────────────────────
#  BATCH PROCESSING
# ─────────────────────────────────────────────────────────────────────────────
def run_batch(files, gemini_key):
    results = []
    for f in files:
        try:
            d1, d2 = parse_two_year_csv(f)
            client  = d2.get('client_name', f.name)
            fy1     = d1.get('fy_year', 'FY2024-25')
            fy2     = d2.get('fy_year', 'FY2025-26')
            fin1    = compute(d1)
            fin2    = compute(d2)
            advice  = (get_gemini_advice(gemini_key, fin1, fin2, fy1, fy2, client)
                       if gemini_key else rule_based_advice(fin2, fy2))
            flags   = audit_risk_flags(fin2, fy2)
            figs_batch = {
                'yoy':     chart_yoy_comparison(fin1, fin2, fy1, fy2),
                'income':  chart_income_breakdown(fin2, fy2),
                'expense': chart_expense_breakdown(fin2, fy2),
                'bracket': chart_tax_bracket(fin2, fy2),
                'rental':  chart_rental_cashflow(fin2, fy2) if fin2['rental_income'] > 0 else None,
                'super':   chart_super_modeller(fin2, fy2) if fin2['super_headroom'] > 500 else None,
                'cgt':     chart_cgt_harvest(fin2, fy2) if (fin2['cgt_gross'] > 0 or fin2['unrealised_loss'] > 0) else None,
            }
            pdf_bytes = generate_pdf(client, fy1, fy2, fin1, fin2, advice, flags, figs_batch)
            results.append({
                'name': client, 'fy': fy2,
                'taxable': fin2['taxable_income'], 'tax': fin2['gross_tax'],
                'refund': fin2['refund'], 'owing': fin2['owing'],
                'pdf': pdf_bytes, 'filename': f.name
            })
            for fig in figs_batch.values():
                if fig: plt.close(fig)
        except Exception as e:
            results.append({'name': f.name, 'error': str(e)})
    return results

# ─────────────────────────────────────────────────────────────────────────────
#  ═══════════════════════════  STEP 1: UPLOAD  ════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
if st.session_state.step == 1:
    st.markdown('<div class="sec-head">📂 Upload Client CSV</div>', unsafe_allow_html=True)

    mode = st.radio("Processing Mode", ["Single Client", "Batch (Multiple Clients)"],
                    horizontal=True)

    if mode == "Single Client":
        uploaded = st.file_uploader("Upload 2-Year Client CSV", type=["csv"],
                                    help="Use the template from the sidebar")
        if uploaded:
            d1, d2 = parse_two_year_csv(uploaded)
            if not d1:
                st.error("Could not parse CSV. Please use the template from the sidebar.")
            else:
                st.session_state.f1   = compute(d1)
                st.session_state.f2   = compute(d2)
                st.session_state.meta = {
                    'client': d2.get('client_name', d1.get('client_name','Client')),
                    'fy1': d1.get('fy_year','FY2024-25'),
                    'fy2': d2.get('fy_year','FY2025-26'),
                }
                st.success(f"✅ Loaded: **{st.session_state.meta['client']}** | "
                           f"{st.session_state.meta['fy1']} & {st.session_state.meta['fy2']}")
                if st.button("▶️ Proceed to Review →", type="primary", use_container_width=True):
                    st.session_state.step = 2
                    st.rerun()

    else:  # Batch
        st.info("Upload multiple 2-year CSVs. Reports will be generated for all clients and packaged as a ZIP.")
        batch_files = st.file_uploader("Upload Client CSVs (multiple)", type=["csv"],
                                       accept_multiple_files=True)
        if batch_files:
            st.write(f"**{len(batch_files)} file(s) selected:**")
            for f in batch_files:
                st.write(f"  • {f.name}")
            if st.button(f"⚡ Process All {len(batch_files)} Clients", type="primary", use_container_width=True):
                with st.spinner(f"Processing {len(batch_files)} clients..."):
                    st.session_state.batch_results = run_batch(batch_files, gemini_key)
                st.session_state.step = 4
                st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
#  ══════════════════════════  STEP 2: REVIEW  ═════════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
elif st.session_state.step == 2:
    f1 = st.session_state.f1
    f2 = st.session_state.f2
    m  = st.session_state.meta
    if not f1 or not f2:
        st.warning("Please go back to Step 1 and upload a CSV.")
        st.stop()

    fy1, fy2, client = m['fy1'], m['fy2'], m['client']

    # ── Exec Summary ────────────────────────────────────────────────────────
    render_exec_summary(f1, f2, fy1, fy2, client,
                        st.session_state.advice or rule_based_advice(f2, fy2))

    # ── Super SG estimation warning ─────────────────────────────────────────
    if f2.get('super_sg_estimated'):
        st.warning(
            f"⚠️ **Super concessional contributions were entered as $0** but this client has "
            f"a salary of **${f2['salary']:,.0f}**. The app has automatically estimated "
            f"employer Super Guarantee (11% = **${f2['super_cc']:,.0f}**) for {fy2}. "
            f"Please verify the actual employer SG and update the CSV with the correct figure "
            f"under `super_contributions_concessional`."
            + (" The concessional cap of $30,000 appears **already exceeded** by employer SG alone — "
               "no further personal deductible contributions are available this year."
               if f2.get('super_cap_exceeded') else "")
        )

    # ATO Verification banner
    if f2.get('ato_verified'):
        st.markdown(
            '<div style="background:#EAF5F0;border-left:4px solid #27AE60;'
            'border-radius:8px;padding:10px 16px;margin-bottom:12px;font-size:13px;">'
            f'<b>&#x2705; ATO-Verified figures active for {fy2}</b> &mdash; '
            'Taxable income and tax payable are from the <b>lodged tax return</b>, '
            'ensuring 100% accuracy. Charts and planning use computed estimates.'
            '</div>', unsafe_allow_html=True)

    # ── KPI Row ──────────────────────────────────────────────────────────────
    st.markdown('<div class="sec-head">&#128202; Key Metrics &mdash; FY26</div>', unsafe_allow_html=True)
    c1,c2,c3,c4,c5 = st.columns(5)
    def kpi(col, label, val, delta=None, delta_good=True, color=DNA_BLUE, prefix="$"):
        dhtml = ""
        if delta is not None:
            dcls = "delta-up" if (delta>=0)==delta_good else "delta-down"
            dhtml = f'<div class="delta {dcls}">{"▲" if delta>=0 else "▼"} ${abs(delta):,.0f} vs {fy1}</div>'
        col.markdown(f"""<div class="kpi-card" style="border-left-color:{color}">
            <div class="label">{label}</div>
            <div class="value" style="color:{color};">{prefix}{val:,.0f}</div>
            {dhtml}</div>""", unsafe_allow_html=True)

    _ato = f2.get('ato_verified', False)
    _dt  = f2['display_taxable']
    _dr  = f2['display_refund']
    _do  = f2['display_owing']
    _badge = ' &#x2713;ATO' if _ato else ''

    with c1: kpi(c1,'Gross Income',      f2['gross_income'],  f2['gross_income']-f1['gross_income'])
    with c2: kpi(c2,'Taxable Income'+_badge, _dt,             _dt - f1['display_taxable'], False)
    with c3: kpi(c3,'Tax & Medicare'+_badge, f2['display_gross_tax'], f2['display_gross_tax']-f1['display_gross_tax'], False, DNA_RED)
    with c4:
        if _dr > 0:
            kpi(c4,'Tax Refund'+_badge, _dr, color=DNA_GREEN)
        else:
            kpi(c4,'Amount Owing'+_badge, _do, color=DNA_RED)
    if f2.get('super_cap_exceeded'):
        c5.markdown(
            f'<div class="kpi-card" style="border-left-color:#E74C3C;">'
            '<div class="label">SUPER CAP EXCEEDED</div>'
            f'<div class="value" style="color:#E74C3C;">-${f2["super_excess"]:,.0f}</div>'
            '<div class="sub">Over $30k cap — review urgently</div>'
            '</div>', unsafe_allow_html=True)
    else:
        kpi(c5,'Super Headroom', f2['super_headroom'], color=DNA_GOLD)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Tabs for analysis sections ────────────────────────────────────────────
    tabs = st.tabs(["📈 Year-on-Year","💰 Income & Expenses","🎯 Tax Bracket",
                    "🏠 Rental Property","💡 Super Modeller","📉 CGT Harvest",
                    "👥 Income Split","⚠️ Audit Risk","📋 YoY Detail Table"])

    with tabs[0]:  # YoY comparison
        st.markdown('<div class="sec-head">Year-on-Year Comparison</div>', unsafe_allow_html=True)
        fig_yoy = chart_yoy_comparison(f1, f2, fy1, fy2)
        st.pyplot(fig_yoy)

    with tabs[1]:  # Income & Expenses
        col_a, col_b = st.columns(2)
        with col_a:
            fig_inc = chart_income_breakdown(f2, fy2)
            if fig_inc: st.pyplot(fig_inc)
        with col_b:
            fig_exp = chart_expense_breakdown(f2, fy2)
            if fig_exp: st.pyplot(fig_exp)
        # Mobile: show fy1 below
        with st.expander(f"View {fy1} Income & Expenses"):
            col_c, col_d = st.columns(2)
            with col_c:
                f1i = chart_income_breakdown(f1, fy1)
                if f1i: st.pyplot(f1i)
            with col_d:
                f1e = chart_expense_breakdown(f1, fy1)
                if f1e: st.pyplot(f1e)

    with tabs[2]:  # Tax bracket
        st.markdown('<div class="sec-head">Tax Bracket Visualiser</div>', unsafe_allow_html=True)
        col_x, col_y = st.columns(2)
        with col_x:
            st.caption(fy1)
            st.pyplot(chart_tax_bracket(f1, fy1))
        with col_y:
            st.caption(fy2)
            st.pyplot(chart_tax_bracket(f2, fy2))

    with tabs[3]:  # Rental
        if f2['rental_income'] > 0:
            st.markdown('<div class="sec-head">Rental Property Cash Flow & Tax Analysis</div>',
                        unsafe_allow_html=True)
            if f2['rental_income'] > 0:
                st.pyplot(chart_rental_cashflow(f2, fy2))
            else:
                st.info('No rental income declared for this client.')
            col_r1,col_r2,col_r3 = st.columns(3)
            col_r1.metric("Cash Flow (pre-depreciation)", f"${f2['rental_cf']:,.0f}")
            col_r2.metric("Net Rental (incl. depreciation)", f"${f2['rental_net']:,.0f}")
            col_r3.metric("Est. Tax Benefit from Loss", f"${f2['rental_tax_benefit']:,.0f}")
        else:
            st.info("No rental income declared for this client.")

    with tabs[4]:  # Super
        st.markdown('<div class="sec-head">Super Contribution Tax Saving Modeller</div>',
                    unsafe_allow_html=True)
        fig_super = chart_super_modeller(f2, fy2) if f2['super_headroom'] > 500 else None
        if fig_super:
            if fig_super:
                st.pyplot(fig_super)
            else:
                st.info('Concessional super cap fully utilised — no additional contribution headroom.')
            st.info(f"💡 Marginal rate {int(f2['marginal_rate']*100)}% vs 15% super tax = "
                    f"**{int((f2['marginal_rate']-0.15)*100)}¢ saved per $1** contributed. "
                    f"Max additional saving: **${f2['max_super_saving']:,.0f}**")
        else:
            st.success("✅ Super concessional cap is already fully utilised.")

    with tabs[5]:  # CGT
        if f2['cgt_gross'] > 0 or f2['unrealised_loss'] > 0:
            st.markdown('<div class="sec-head">CGT Harvest Opportunity Analysis</div>',
                        unsafe_allow_html=True)
            fig_cgt = chart_cgt_harvest(f2, fy2) if (f2['cgt_gross'] > 0 or f2['unrealised_loss'] > 0) else None
            if fig_cgt: st.pyplot(fig_cgt)
        else:
            st.info("No capital gains or unrealised losses declared.")

    with tabs[6]:  # Income split
        if f2['spouse_inc'] > 0:
            st.markdown('<div class="sec-head">Household Income Split Analysis</div>',
                        unsafe_allow_html=True)
            fig_sp = chart_spouse_split(f2, fy2)
            if fig_sp: st.pyplot(fig_sp)
            if f2['spouse_split_saving'] > 0:
                st.success(f"💡 Potential household tax saving from income equalisation: "
                           f"**${f2['spouse_split_saving']:,.0f}**")
        else:
            st.info("No spouse income declared. Enter spouse_income in the CSV to enable this analysis.")

    with tabs[7]:  # Audit risk
        st.markdown('<div class="sec-head">⚠️ ATO Audit Risk Indicators</div>', unsafe_allow_html=True)
        flags = audit_risk_flags(f2, fy2)
        for level, msg in flags:
            col_map = {"High":"risk-high","Medium":"risk-medium","Low":"risk-low"}
            st.markdown(f'<div class="{col_map.get(level,"risk-low")}">'
                        f'<b>{level} Risk:</b> {msg}</div>', unsafe_allow_html=True)

    with tabs[8]:  # YoY detail table
        st.markdown('<div class="sec-head">Year-on-Year Detail Table</div>', unsafe_allow_html=True)
        render_yoy_table(f1, f2, fy1, fy2)

    st.markdown("<br>", unsafe_allow_html=True)
    col_nav1, col_nav2 = st.columns(2)
    with col_nav1:
        if st.button("← Back to Upload", use_container_width=True):
            st.session_state.step = 1; st.rerun()
    with col_nav2:
        if st.button("▶️ Proceed to AI Advice →", type="primary", use_container_width=True):
            st.session_state.step = 3; st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
#  ══════════════════════════  STEP 3: AI ADVICE  ══════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
elif st.session_state.step == 3:
    f1 = st.session_state.f1
    f2 = st.session_state.f2
    m  = st.session_state.meta
    fy1, fy2, client = m['fy1'], m['fy2'], m['client']

    st.markdown('<div class="sec-head">🤖 AI-Powered Tax Optimisation Advice</div>',
                unsafe_allow_html=True)
    st.caption(f"Client: **{client}** | Two-year analysis: {fy1} & {fy2}")

    if not st.session_state.advice:
        if gemini_key:
            st.markdown(
                f'''<div style="background:#EAF5F5;border-left:4px solid #006F73;border-radius:8px;
                padding:10px 14px;margin-bottom:12px;font-size:13px;">
                🤖 <b>Gemini AI Mode</b> — Generating personalised advice based on
                {client}'s two-year financial profile...</div>''',
                unsafe_allow_html=True)
            with st.spinner("🔮 Gemini AI is analysing your data..."):
                st.session_state.advice = get_gemini_advice(
                    gemini_key, f1, f2, fy1, fy2, client)
        else:
            st.markdown(
                '''<div style="background:#FFF8EC;border-left:4px solid #C9A84C;border-radius:8px;
                padding:10px 14px;margin-bottom:12px;font-size:13px;">
                📋 <b>Rule-Based Mode</b> — AI toggle is OFF or no API key provided.
                Showing DNA's built-in ATO tax optimisation advice.
                <br><span style="font-size:11px;color:#888;">
                Enable the AI toggle in the sidebar to use Gemini AI.</span></div>''',
                unsafe_allow_html=True)
            st.session_state.advice = rule_based_advice(f2, fy2)

    advice = st.session_state.advice
    pri_col = {"High": DNA_RED, "Medium": DNA_GOLD, "Low": DNA_GREEN}

    for i, tip in enumerate(advice, 1):
        pri = tip.get('priority','Medium')
        pc  = pri_col.get(pri, DNA_GOLD)
        st.markdown(f"""
        <div class="adv-card">
          <div style="display:flex;align-items:center;gap:10px;margin-bottom:6px;">
            <span style="font-size:20px;font-weight:800;color:{pc};">{i}</span>
            <span style="background:{pc};color:white;font-size:11px;font-weight:700;
                   padding:2px 8px;border-radius:10px;">{pri} Priority</span>
            <span class="atitle">{tip.get('title','')}</span>
          </div>
          <div class="abody">{tip.get('detail','')}</div>
          <div style="color:{DNA_GREEN};font-weight:700;font-size:13px;margin-top:6px;">
            💰 {tip.get('estimated_saving','N/A')}
          </div>
        </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    if st.button("🔄 Regenerate Advice", use_container_width=False):
        st.session_state.advice = []
        st.rerun()

    col_nav1, col_nav2 = st.columns(2)
    with col_nav1:
        if st.button("← Back to Review", use_container_width=True):
            st.session_state.step = 2; st.rerun()
    with col_nav2:
        if st.button("▶️ Proceed to Download →", type="primary", use_container_width=True):
            st.session_state.step = 4; st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
#  ══════════════════════════  STEP 4: DOWNLOAD  ═══════════════════════════════
# ─────────────────────────────────────────────────────────────────────────────
elif st.session_state.step == 4:

    # ── Batch download ────────────────────────────────────────────────────────
    if st.session_state.batch_results:
        st.markdown('<div class="sec-head">⚡ Batch Processing Results</div>', unsafe_allow_html=True)
        results = st.session_state.batch_results
        ok  = [r for r in results if 'pdf' in r]
        err = [r for r in results if 'error' in r]

        if ok:
            summary_rows = [{"Client": r['name'], "FY": r['fy'],
                             "Taxable Income": f"${r['taxable']:,.0f}",
                             "Tax": f"${r['tax']:,.0f}",
                             "Refund/(Owing)": f"${r['refund']:,.0f}" if r['refund']>0 else f"(${r['owing']:,.0f})"}
                            for r in ok]
            st.dataframe(pd.DataFrame(summary_rows), use_container_width=True, hide_index=True)

            # Build ZIP
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                for r in ok:
                    zf.writestr(f"DNA_TaxReport_{r['name'].replace(' ','_')}.pdf", r['pdf'])
            zip_buf.seek(0)
            st.download_button("📦 Download All Reports (ZIP)",
                               zip_buf.getvalue(),
                               f"DNA_Batch_Reports_{datetime.now().strftime('%Y%m%d')}.zip",
                               "application/zip", use_container_width=True)

        if err:
            st.error("Errors in batch:")
            for e in err:
                st.write(f"• {e['name']}: {e['error']}")

        if st.button("← Start New Batch", use_container_width=True):
            st.session_state.batch_results = []
            st.session_state.step = 1; st.rerun()

    # ── Single client download ─────────────────────────────────────────────────
    else:
        f1 = st.session_state.f1
        f2 = st.session_state.f2
        m  = st.session_state.meta
        if not f1 or not f2:
            st.warning("No data loaded. Return to Step 1.")
            st.stop()

        fy1, fy2, client = m['fy1'], m['fy2'], m['client']
        advice = st.session_state.advice or rule_based_advice(f2, fy2)
        flags  = audit_risk_flags(f2, fy2)

        st.markdown('<div class="sec-head">⬇️ Download Reports</div>', unsafe_allow_html=True)
        st.caption(f"Client: **{client}** | {fy1} & {fy2}")

        with st.spinner("Building charts & reports..."):
            figs = {
                'yoy':     chart_yoy_comparison(f1, f2, fy1, fy2),
                'income':  chart_income_breakdown(f2, fy2),
                'expense': chart_expense_breakdown(f2, fy2),
                'bracket': chart_tax_bracket(f2, fy2),
                'rental':  chart_rental_cashflow(f2, fy2) if f2['rental_income'] > 0 else None,
                'super':   chart_super_modeller(f2, fy2),
                'cgt':     chart_cgt_harvest(f2, fy2),
                'spouse':  chart_spouse_split(f2, fy2) if f2['spouse_inc'] > 0 else None,
            }
            pdf_bytes  = generate_pdf(client, fy1, fy2, f1, f2, advice, flags, figs)
            docx_bytes = generate_docx(client, fy1, fy2, f1, f2, advice, flags, figs)
            for fig in figs.values():
                if fig: plt.close(fig)

        name_slug = client.replace(' ', '_')
        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button("📄 Download PDF Report", pdf_bytes,
                               f"DNA_Tax_Report_{name_slug}_{fy2}.pdf",
                               "application/pdf", use_container_width=True)
        with col2:
            st.download_button("📝 Download Word Report", docx_bytes,
                               f"DNA_Tax_Report_{name_slug}_{fy2}.docx",
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               use_container_width=True)
        with col3:
            summary_csv = pd.DataFrame({
                "Metric":["Gross Income","Taxable Income","Tax Payable","Medicare",
                          "Offsets","PAYG Credits","Refund/(Owing)","Effective Rate"],
                fy1:[f"${f1['gross_income']:,.0f}",f"${f1['taxable_income']:,.0f}",
                     f"${f1['income_tax']:,.0f}",f"${f1['medicare']:,.0f}",
                     "${:,.0f}".format(f1['lito']+f1['seniors_off']), "${:,.0f}".format(f1['credits']),
                     f"${f1['refund']:,.0f}" if f1['refund']>0 else f"(${f1['owing']:,.0f})",
                     f"{f1['eff_rate']}%"],
                fy2:[f"${f2['gross_income']:,.0f}",f"${f2['taxable_income']:,.0f}",
                     f"${f2['income_tax']:,.0f}",f"${f2['medicare']:,.0f}",
                     "${:,.0f}".format(f2['lito']+f2['seniors_off']), "${:,.0f}".format(f2['credits']),
                     f"${f2['refund']:,.0f}" if f2['refund']>0 else f"(${f2['owing']:,.0f})",
                     f"{f2['eff_rate']}%"],
            }).to_csv(index=False)
            st.download_button("📊 Download Summary CSV", summary_csv,
                               f"DNA_Tax_Summary_{name_slug}.csv",
                               "text/csv", use_container_width=True)

        # Mini re-cap
        st.markdown("<br>", unsafe_allow_html=True)
        render_exec_summary(f1, f2, fy1, fy2, client, advice)

        if st.button("← Back to AI Advice", use_container_width=False):
            st.session_state.step = 3; st.rerun()
        if st.button("🔄 Start New Client", use_container_width=False):
            for k in ['f1','f2','meta','advice','batch_results']:
                st.session_state[k] = None if k in ('f1','f2','meta') else []
            st.session_state.meta = {}
            st.session_state.step = 1; st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
#  FOOTER
# ─────────────────────────────────────────────────────────────────────────────
st.markdown(f"""
<div style="background:{DNA_BLUE};color:white;padding:14px 24px;border-radius:10px;
            font-size:11px;text-align:center;margin-top:32px;opacity:.92;">
  📋 DNA Accountants Pty Ltd &nbsp;|&nbsp; Chartered Accountants &nbsp;|&nbsp;
  110 Pitt Street, Sydney NSW 2000 &nbsp;|&nbsp; (02) 9064 4400 &nbsp;|&nbsp; This report is for guidance only. &nbsp;|&nbsp;
  v3.0 &nbsp;|&nbsp; {datetime.now().strftime('%d %B %Y')}
</div>
""", unsafe_allow_html=True)
